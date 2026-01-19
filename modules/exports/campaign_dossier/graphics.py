from __future__ import annotations

import os

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

MIN_FONT_SIZE_PT = 1

COVER_STAMP_ASSET = os.path.join("images", "postage-stamp.png")
SECTION_DIVIDER_ASSET = os.path.join("images", "doodle_symbol.png")


def _repo_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))


def _asset_path(relative_path: str) -> str:
    return os.path.join(_repo_root(), "assets", relative_path)


def _apply_font_size(run, size_pt: float) -> None:
    run.font.size = Pt(max(size_pt, MIN_FONT_SIZE_PT))


def _add_picture(document, asset_path: str, width, alignment) -> bool:
    if not os.path.exists(asset_path):
        return False
    document.add_picture(asset_path, width=width)
    document.paragraphs[-1].alignment = alignment
    return True


def _apply_cell_shading(cell, color_rgb) -> None:
    shading = OxmlElement("w:shd")
    fill = "{:02X}{:02X}{:02X}".format(color_rgb[0], color_rgb[1], color_rgb[2])
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_confidential_header(document, theme_meta: dict) -> None:
    header = document.sections[0].header
    paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.color.rgb = theme_meta["accent"]
    _apply_font_size(run, 8)


def add_cover_page(document, theme_meta: dict) -> None:
    section = document.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    stamp_path = _asset_path(COVER_STAMP_ASSET)
    _add_picture(
        document,
        stamp_path,
        width=int(available_width * 0.3),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Campaign Dossier")
    title_run.bold = True
    title_run.font.name = theme_meta["heading_font"]
    _apply_font_size(title_run, 28)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Theme: {theme_meta.get('name', 'Dossier')}")
    subtitle_run.font.name = theme_meta["font"]
    subtitle_run.font.color.rgb = theme_meta["accent"]
    _apply_font_size(subtitle_run, 12)

    confidential = document.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    confidential_run = confidential.add_run("CONFIDENTIAL")
    confidential_run.bold = True
    confidential_run.font.name = theme_meta["heading_font"]
    confidential_run.font.color.rgb = theme_meta["accent"]
    _apply_font_size(confidential_run, 14)

    bar_table = document.add_table(rows=1, cols=1)
    bar_table.allow_autofit = False
    bar_table.columns[0].width = available_width
    bar_table.rows[0].height = Pt(4)
    bar_cell = bar_table.cell(0, 0)
    bar_cell.text = ""
    _apply_cell_shading(bar_cell, theme_meta["accent"])

    document.add_page_break()


def add_section_divider(document) -> None:
    section = document.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    divider_path = _asset_path(SECTION_DIVIDER_ASSET)
    _add_picture(
        document,
        divider_path,
        width=int(available_width * 0.2),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
