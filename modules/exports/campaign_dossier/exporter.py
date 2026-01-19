from __future__ import annotations

import importlib.util
import os
from dataclasses import dataclass
from typing import Iterable

from docx import Document

from modules.generic.generic_model_wrapper import GenericModelWrapper
from modules.helpers.logging_helper import log_warning
from modules.helpers.template_loader import load_entity_definitions, load_template, list_known_entities
from modules.helpers.text_helpers import deserialize_possible_json, normalize_rtf_json
from modules.helpers.theme_manager import get_theme

from modules.exports.campaign_dossier.graphics import (
    add_confidential_header,
    add_cover_page,
    add_section_divider,
)
from modules.exports.campaign_dossier.layouts import apply_layout, format_entity_label
from modules.exports.campaign_dossier.templates import apply_dossier_theme


@dataclass(frozen=True)
class DossierExportOptions:
    layout_key: str
    pagination_mode: str
    include_toc: bool
    include_branding: bool
    output_mode: str
    output_format: str
    output_target: str


@dataclass(frozen=True)
class EntityGroup:
    slug: str
    label: str
    items: list[dict]


def _docx2pdf_available() -> bool:
    return importlib.util.find_spec("docx2pdf") is not None


def _normalize_name(record: dict) -> str:
    name = record.get("Name") or record.get("Title") or "Unnamed"
    return str(name).strip() or "Unnamed"


def _sanitize_filename(value: str) -> str:
    safe = "".join(ch if ch.isalnum() or ch in "-_ " else "_" for ch in value)
    safe = "_".join(safe.split())
    return safe or "entity"


def _normalize_rtf_payload(value) -> dict:
    parsed = deserialize_possible_json(value)
    if isinstance(parsed, dict):
        normalized = normalize_rtf_json(parsed)
        return {
            "text": str(normalized.get("text", "")),
            "formatting": normalized.get("formatting", {}) or {},
        }
    return {"text": str(parsed or ""), "formatting": {}}


def _format_rtf_runs(text: str, formatting: dict) -> list[tuple[str, dict]]:
    if not formatting:
        return [(text, {})]

    opens: dict[int, list[str]] = {}
    closes: dict[int, list[str]] = {}
    for tag in ("bold", "italic", "underline"):
        ranges = formatting.get(tag, []) or []
        for start, end in ranges:
            try:
                start_index = int(start)
                end_index = int(end)
            except (TypeError, ValueError):
                continue
            start_index = max(0, min(start_index, len(text)))
            end_index = max(start_index, min(end_index, len(text)))
            if start_index == end_index:
                continue
            opens.setdefault(start_index, []).append(tag)
            closes.setdefault(end_index, []).append(tag)

    segments: list[tuple[str, dict]] = []
    active: set[str] = set()
    current: list[str] = []

    for index, char in enumerate(text):
        if index in closes or index in opens:
            if current:
                segments.append(
                    (
                        "".join(current),
                        {
                            "bold": "bold" in active,
                            "italic": "italic" in active,
                            "underline": "underline" in active,
                        },
                    )
                )
                current = []
            for tag in closes.get(index, []):
                active.discard(tag)
            for tag in opens.get(index, []):
                active.add(tag)
        current.append(char)

    if current:
        segments.append(
            (
                "".join(current),
                {
                    "bold": "bold" in active,
                    "italic": "italic" in active,
                    "underline": "underline" in active,
                },
            )
        )
    return segments


def _format_field_value(value) -> list[tuple[str, dict]]:
    if value is None:
        return []
    value = deserialize_possible_json(value)
    if isinstance(value, dict):
        if "text" in value or "formatting" in value:
            payload = _normalize_rtf_payload(value)
            return _format_rtf_runs(payload.get("text", ""), payload.get("formatting", {}))
        return [(str(value), {})]
    if isinstance(value, list):
        segments: list[tuple[str, dict]] = []
        for item in value:
            item_segments = _format_field_value(deserialize_possible_json(item))
            if item_segments:
                if segments:
                    segments.append(("\n", {}))
                segments.extend(item_segments)
        return segments
    return [(str(value), {})]


def _add_field_paragraph(document, label: str, value) -> None:
    text_parts = _format_field_value(value)
    if not text_parts:
        return
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    for text, formatting in text_parts:
        if not text:
            continue
        run = paragraph.add_run(text)
        if formatting.get("bold"):
            run.bold = True
        if formatting.get("italic"):
            run.italic = True
        if formatting.get("underline"):
            run.underline = True


def _add_entity_section(document, record: dict, fields: Iterable[dict]) -> None:
    for index, field in enumerate(fields):
        if index == 0:
            continue
        name = field.get("name")
        if not name:
            continue
        value = record.get(name)
        if value in (None, "", [], {}):
            continue
        _add_field_paragraph(document, str(name), value)


def _collect_entities() -> list[EntityGroup]:
    definitions = load_entity_definitions()
    groups: list[EntityGroup] = []
    for slug in list_known_entities():
        label = definitions.get(slug, {}).get("label") or slug.replace("_", " ").title()
        wrapper = GenericModelWrapper(slug)
        items = wrapper.load_items()
        if not items:
            continue
        items_sorted = sorted(items, key=_normalize_name)
        groups.append(EntityGroup(slug=slug, label=label, items=items_sorted))
    return groups


def _add_table_of_contents(document, groups: list[EntityGroup], preset) -> None:
    document.add_heading("Table of Contents", level=1)
    for group in groups:
        document.add_paragraph(group.label, style="List Bullet")
        for record in group.items:
            document.add_paragraph(
                format_entity_label(preset, group.label, _normalize_name(record)),
                style="List Bullet 2",
            )
    document.add_page_break()


def _build_document(groups: list[EntityGroup], options: DossierExportOptions) -> Document:
    theme_key = get_theme()
    document = Document()
    preset = apply_layout(document, options.layout_key, options.include_branding)
    theme_meta = apply_dossier_theme(document, theme_key)

    add_confidential_header(document, theme_meta)
    add_cover_page(document, theme_meta)

    if options.include_toc:
        _add_table_of_contents(document, groups, preset)

    for group_index, group in enumerate(groups):
        if group_index > 0:
            add_section_divider(document)
        document.add_heading(group.label, level=2)
        template = load_template(group.slug)
        fields = template.get("fields", [])

        for record_index, record in enumerate(group.items):
            document.add_heading(
                format_entity_label(preset, group.label, _normalize_name(record)),
                level=3,
            )
            _add_entity_section(document, record, fields)
            is_last_record = record_index == len(group.items) - 1
            is_last_group = group_index == len(groups) - 1
            is_last_entity = is_last_record and is_last_group
            if options.pagination_mode == "entity" and not is_last_entity:
                document.add_page_break()
        if options.pagination_mode == "section" and not is_last_group:
            document.add_page_break()

    return document


def _save_docx(document: Document, path: str) -> None:
    directory = os.path.dirname(path)
    if directory:
        os.makedirs(directory, exist_ok=True)
    document.save(path)


def _convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    if not _docx2pdf_available():
        return False
    from docx2pdf import convert
    try:
        convert(docx_path, pdf_path)
    except Exception as exc:
        log_warning(
            f"PDF conversion failed: {exc}",
            func_name="modules.exports.campaign_dossier.exporter._convert_docx_to_pdf",
        )
        return False
    return os.path.exists(pdf_path)


def _resolve_output_path(target: str, name: str, output_format: str) -> str:
    ext = ".pdf" if output_format == "pdf" else ".docx"
    return os.path.join(target, f"{_sanitize_filename(name)}{ext}")


def export_campaign_dossier(options: DossierExportOptions) -> list[str]:
    groups = _collect_entities()
    if not groups:
        return []

    output_paths: list[str] = []

    if options.output_mode == "folder":
        for group in groups:
            group_dir = os.path.join(options.output_target, _sanitize_filename(group.label))
            os.makedirs(group_dir, exist_ok=True)
            template = load_template(group.slug)
            fields = template.get("fields", [])
            for record in group.items:
                document = Document()
                preset = apply_layout(document, options.layout_key, options.include_branding)
                theme_meta = apply_dossier_theme(document, get_theme())
                add_confidential_header(document, theme_meta)
                document.add_heading(group.label, level=2)
                document.add_heading(
                    format_entity_label(preset, group.label, _normalize_name(record)),
                    level=3,
                )
                _add_entity_section(document, record, fields)
                docx_path = _resolve_output_path(group_dir, _normalize_name(record), "docx")
                _save_docx(document, docx_path)
                if options.output_format == "pdf":
                    pdf_path = _resolve_output_path(group_dir, _normalize_name(record), "pdf")
                    if _convert_docx_to_pdf(docx_path, pdf_path):
                        output_paths.append(pdf_path)
                    else:
                        output_paths.append(docx_path)
                else:
                    output_paths.append(docx_path)
        return output_paths

    document = _build_document(groups, options)
    output_path = options.output_target
    if options.output_format == "pdf":
        docx_path = os.path.splitext(output_path)[0] + ".docx"
        _save_docx(document, docx_path)
        if _convert_docx_to_pdf(docx_path, output_path):
            output_paths.append(output_path)
            return output_paths
        log_warning(
            "PDF conversion unavailable; leaving DOCX export in place.",
            func_name="modules.exports.campaign_dossier.exporter.export_campaign_dossier",
        )
        output_paths.append(docx_path)
        return output_paths

    _save_docx(document, output_path)
    output_paths.append(output_path)
    return output_paths
