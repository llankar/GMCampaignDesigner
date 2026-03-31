"""Utilities for session brief exporter."""

from __future__ import annotations

import os
from pathlib import Path

from typing import TYPE_CHECKING

from modules.helpers.logging_helper import log_warning

if TYPE_CHECKING:
    from docx import Document


def _sanitize_filename(value: str) -> str:
    """Internal helper for sanitize filename."""
    safe = "".join(ch if ch.isalnum() or ch in "-_ " else "_" for ch in value)
    safe = "_".join(safe.split())
    return safe or "session_brief"


def _build_markdown(
    *,
    campaign_name: str,
    summary: str,
    active_arcs: list[str],
    arc_details: list[str],
    dashboard_fields: list[str],
    gm_priority_notes: list[str],
) -> str:
    """Build markdown."""
    lines: list[str] = [
        f"# Session brief — {campaign_name}",
        "",
        "## Summary",
        summary.strip() or "No summary available.",
        "",
        "## Active arcs",
    ]

    if active_arcs:
        lines.extend([f"- {arc}" for arc in active_arcs])
    else:
        lines.append("- No active arc.")

    lines.extend(["", "## Arc details"])
    if arc_details:
        lines.extend([f"- {arc}" for arc in arc_details])
    else:
        lines.append("- No arc details.")

    lines.extend(["", "## Campaign"])
    if dashboard_fields:
        lines.extend([f"- {field}" for field in dashboard_fields])
    else:
        lines.append("- No campaign field")

    lines.extend(["", "## Priority GM notes"])
    if gm_priority_notes:
        lines.extend([f"- {note}" for note in gm_priority_notes])
    else:
        lines.append("- No priority note.")

    lines.append("")
    return "\n".join(lines)


def _convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Internal helper for convert docx to PDF."""
    try:
        from docx2pdf import convert
    except Exception:
        return False

    try:
        convert(docx_path, pdf_path)
    except Exception as exc:
        log_warning(
            f"Session brief PDF conversion failed: {exc}",
            func_name="modules.exports.session_brief.exporter._convert_docx_to_pdf",
        )
        return False
    return os.path.exists(pdf_path)


def _build_docx_document(
    *,
    campaign_name: str,
    summary: str,
    active_arcs: list[str],
    arc_details: list[str],
    dashboard_fields: list[str],
    gm_priority_notes: list[str],
):
    """Build docx document."""
    from docx import Document

    document = Document()
    document.add_heading(f"Session brief — {campaign_name}", level=1)

    def add_section(title: str, values: list[str] | None = None, paragraph_text: str | None = None) -> None:
        """Handle add section."""
        document.add_heading(title, level=2)
        if paragraph_text is not None:
            document.add_paragraph(paragraph_text)
            return
        for value in values or []:
            document.add_paragraph(value, style="List Bullet")

    add_section("Summary", paragraph_text=summary.strip() or "No summary available.")
    add_section("Active arcs", values=active_arcs or ["No active arc."])
    add_section("Arc details", values=arc_details or ["No arc details."])
    add_section("Dashboard fields", values=dashboard_fields or ["No dashboard field."])
    add_section("Priority GM notes", values=gm_priority_notes or ["No priority note."])

    return document


def export_session_brief(
    *,
    campaign_name: str,
    summary: str,
    active_arcs: list[str],
    arc_details: list[str],
    dashboard_fields: list[str],
    gm_priority_notes: list[str],
    output_format: str,
    output_path: str,
) -> str:
    """Export a session brief to markdown, DOCX, or lightweight PDF.

    Returns the exported file path.
    """

    export_format = (output_format or "markdown").strip().lower()
    target_path = Path(output_path)

    if export_format == "markdown":
        # Handle the branch where export_format == 'markdown'.
        if target_path.suffix.lower() != ".md":
            target_path = target_path.with_suffix(".md")
        payload = _build_markdown(
            campaign_name=campaign_name,
            summary=summary,
            active_arcs=active_arcs,
            arc_details=arc_details,
            dashboard_fields=dashboard_fields,
            gm_priority_notes=gm_priority_notes,
        )
        target_path.write_text(payload, encoding="utf-8")
        return str(target_path)

    if export_format in {"docx", "pdf"}:
        # Handle the branch where export format is in {'docx', 'pdf'}.
        docx_path = target_path.with_suffix(".docx")
        try:
            # Keep session brief resilient if this step fails.
            document = _build_docx_document(
                campaign_name=campaign_name,
                summary=summary,
                active_arcs=active_arcs,
                arc_details=arc_details,
                dashboard_fields=dashboard_fields,
                gm_priority_notes=gm_priority_notes,
            )
            document.save(str(docx_path))
        except Exception as exc:
            log_warning(
                f"Session brief DOCX generation failed: {exc}",
                func_name="modules.exports.session_brief.exporter.export_session_brief",
            )
            markdown_fallback_path = target_path.with_suffix(".md")
            markdown_fallback_path.write_text(
                _build_markdown(
                    campaign_name=campaign_name,
                    summary=summary,
                    active_arcs=active_arcs,
                    arc_details=arc_details,
                    dashboard_fields=dashboard_fields,
                    gm_priority_notes=gm_priority_notes,
                ),
                encoding="utf-8",
            )
            return str(markdown_fallback_path)

        if export_format == "docx":
            return str(docx_path)

        pdf_path = target_path.with_suffix(".pdf")
        if _convert_docx_to_pdf(str(docx_path), str(pdf_path)):
            # Handle the branch where _convert_docx_to_pdf(str(docx_path), str(pdf_path)).
            try:
                docx_path.unlink(missing_ok=True)
            except Exception:
                pass
            return str(pdf_path)

        return str(docx_path)

    suggested_name = _sanitize_filename(campaign_name)
    raise ValueError(f"Unsupported format '{output_format}' for {suggested_name}")
