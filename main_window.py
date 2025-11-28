import os
import sys
import json
import sqlite3
import subprocess
import time
import requests
import shutil
import re
import unicodedata
import threading
from contextlib import contextmanager
from pathlib import Path
import tkinter as tk
from tkinter import (
    filedialog,
    messagebox,
    Toplevel,
    Listbox,
    MULTIPLE,
    PhotoImage,
    simpledialog,
)
from tkinter import ttk

import customtkinter as ctk
from PIL import Image
from docx import Document

# Modular helper imports
from modules.helpers.window_helper import position_window_at_top
from modules.helpers import theme_manager
from modules.helpers.template_loader import (
    load_template,
    load_entity_definitions,
    build_entity_wrappers,
)
from modules.helpers.config_helper import ConfigHelper
from modules.helpers.backup_helper import (
    BackupError,
    ManifestError,
    create_backup_archive,
    read_backup_manifest,
    restore_backup_archive,
)
from modules.helpers.swarmui_helper import get_available_models
from modules.helpers.logging_helper import (
    initialize_logging,
    log_module_import,
    log_debug,
    log_exception,
    log_info,
    log_methods,
    log_warning,
)
from modules.helpers import system_config, update_helper
from modules.helpers.system_config import register_system_change_listener
from modules.ui.tooltip import ToolTip
from modules.ui.icon_button import create_icon_button
from modules.ui.portrait_importer import PortraitImporter
from modules.ui.system_selector_dialog import CampaignSystemSelectorDialog
from modules.ui.database_manager_dialog import DatabaseManagerDialog

from modules.generic.generic_list_view import GenericListView
from modules.generic.generic_model_wrapper import GenericModelWrapper
from modules.scenarios.gm_screen_view import GMScreenView
from modules.scenarios.gm_layout_manager import GMScreenLayoutManager
from modules.npcs.npc_graph_editor import NPCGraphEditor
from modules.pcs.pc_graph_editor import PCGraphEditor
from modules.scenarios.scenario_graph_editor import ScenarioGraphEditor
from modules.scenarios.scenario_importer import ScenarioImportWindow
from modules.objects.object_importer import ObjectImportWindow
from modules.scenarios.scenario_generator_view import ScenarioGeneratorView
from modules.scenarios.scenario_builder_wizard import ScenarioBuilderWizard
from modules.scenarios.random_tables_editor import RandomTableEditorDialog
from modules.generic.export_for_foundry import preview_and_export_foundry
from modules.generic.cross_campaign_asset_library import CrossCampaignAssetLibraryWindow
from modules.helpers import text_helpers, dice_markup
from db.db import initialize_db, ensure_entity_schema
from modules.factions.faction_graph_editor import FactionGraphEditor
from modules.whiteboard.controllers.whiteboard_controller import WhiteboardController
from modules.pcs.display_pcs import display_pcs_in_banner
from modules.generic.generic_list_selection_view import GenericListSelectionView
from modules.maps.controllers.display_map_controller import DisplayMapController
from modules.maps.world_map_view import WorldMapWindow
from modules.generic.custom_fields_editor import CustomFieldsEditor
from modules.generic.new_entity_type_dialog import NewEntityTypeDialog


from modules.dice.dice_roller_window import DiceRollerWindow
from modules.dice.dice_bar_window import DiceBarWindow
from modules.audio.audio_bar_window import AudioBarWindow
from modules.audio.audio_controller import get_audio_controller
from modules.audio.sound_manager_window import SoundManagerWindow

initialize_logging()
log_module_import(__name__)

# Set up CustomTkinter appearance
ctk.set_appearance_mode("Dark")
# Apply configured theme palette (default/medieval/sf)
theme_manager.apply_theme(theme_manager.get_theme())

# Global process variable for SwarmUI
SWARMUI_PROCESS = None

#logging.basicConfig(level=logging.DEBUG, format='%(levelname)s: %(message)s')


@log_methods
class MainWindow(ctk.CTk):
    def __init__(self):
        super().__init__()

        log_info("Initializing MainWindow", func_name="main_window.MainWindow.__init__")

        self.title("GMCampaignDesigner")
        self.geometry("1920x980")
        self.minsize(1920, 980)
        self.attributes("-fullscreen", True)
        self.current_open_view   = None
        self.current_open_entity = None    # ← initialize here to avoid AttributeError
        initialize_db()
        log_info("Database initialization complete", func_name="main_window.MainWindow.__init__")
        position_window_at_top(self)
        self.set_window_icon()
        self.create_layout()
        self.sidebar_default_width = 220
        self._sidebar_collapsed = False
        self._sidebar_animating = False
        self._sidebar_animation_job = None
        self._sidebar_pack_kwargs = None
        self.entity_definitions = load_entity_definitions()
        self.entity_wrappers = {}
        self.load_icons()
        self.create_sidebar()
        self.create_content_area()
        self.create_exit_button()
        self.load_model_config()
        self.init_wrappers()
        self._normalize_entity_media_paths()
        self.current_gm_view = None
        self._gm_mode = False
        self.dice_roller_window = None
        self.dice_bar_window = None
        self.audio_controller = get_audio_controller()
        self.sound_manager_window = None
        self.audio_bar_window = None
        self.portrait_importer = PortraitImporter(self)
        self._asset_library_window = None
        self._busy_modal = None
        self._system_selector_dialog = None
        self._database_manager_dialog = None
        self._update_thread = None
        self.whiteboard_controller = None
        root = self.winfo_toplevel()
        root.bind_all("<Control-f>", self._on_ctrl_f)

        self._system_listener_unsub = register_system_change_listener(self._on_system_changed)
        # Rebuild colorized UI bits when theme changes
        self._theme_listener_unsub = theme_manager.register_theme_change_listener(self._on_theme_changed)

        self.after(200, self.open_dice_bar)
        self.after(400, self.open_audio_bar)
        self.after(600, lambda: self._queue_update_check(force=True))
        self.after(800, self._auto_open_gm_screen_if_available)

    def open_ai_settings(self):
        log_info("Opening AI settings dialog", func_name="main_window.MainWindow.open_ai_settings")
        top = ctk.CTkToplevel(self)
        top.title("AI Settings")
        top.geometry("520x360")
        top.lift(); top.focus_force(); top.attributes("-topmost", True); top.after_idle(lambda: top.attributes("-topmost", False))

        # Current config values
        base_url = ConfigHelper.get("AI", "base_url", fallback="http://localhost:8080") or ""
        model = ConfigHelper.get("AI", "model", fallback="gpt-oss") or ""
        temperature = ConfigHelper.get("AI", "temperature", fallback="0.7") or "0.7"
        max_tokens = ConfigHelper.get("AI", "max_tokens", fallback="512") or "512"
        api_key = ConfigHelper.get("AI", "api_key", fallback="") or ""

        # Vars
        v_base = ctk.StringVar(value=base_url)
        v_model = ctk.StringVar(value=model)
        v_temp = ctk.StringVar(value=str(temperature))
        v_max = ctk.StringVar(value=str(max_tokens))
        v_key = ctk.StringVar(value=api_key)

        form = ctk.CTkFrame(top)
        form.pack(fill="both", expand=True, padx=12, pady=12)

        def row(label, widget):
            r = ctk.CTkFrame(form)
            r.pack(fill="x", pady=6)
            ctk.CTkLabel(r, text=label, width=140, anchor="w").pack(side="left")
            widget.pack(side="left", fill="x", expand=True)

        row("Base URL", ctk.CTkEntry(form, textvariable=v_base, placeholder_text="http://localhost:8080"))
        row("Model", ctk.CTkEntry(form, textvariable=v_model, placeholder_text="gpt-oss"))
        row("Temperature", ctk.CTkEntry(form, textvariable=v_temp))
        row("Max Tokens", ctk.CTkEntry(form, textvariable=v_max))
        row("API Key", ctk.CTkEntry(form, textvariable=v_key))

        btns = ctk.CTkFrame(top)
        btns.pack(fill="x", padx=12, pady=(0,12))

        def save():
            try:
                # Basic normalization
                _ = float(v_temp.get())
                _ = int(v_max.get())
            except Exception:
                log_warning("Rejected AI settings save due to invalid numeric values",
                            func_name="main_window.MainWindow.open_ai_settings.save")
                messagebox.showerror("Invalid Values", "Temperature must be a float and Max Tokens an integer.")
                return
            ConfigHelper.set("AI", "base_url", v_base.get().strip())
            ConfigHelper.set("AI", "model", v_model.get().strip())
            ConfigHelper.set("AI", "temperature", v_temp.get().strip())
            ConfigHelper.set("AI", "max_tokens", v_max.get().strip())
            ConfigHelper.set("AI", "api_key", v_key.get())
            log_info(
                "AI settings saved",
                func_name="main_window.MainWindow.open_ai_settings.save",
            )
            messagebox.showinfo("Saved", "AI settings saved.")

        def reset_defaults():
            v_base.set("http://localhost:8080")
            v_model.set("gpt-oss")
            v_temp.set("0.7")
            v_max.set("512")
            v_key.set("")
            log_info(
                "AI settings reset to defaults",
                func_name="main_window.MainWindow.open_ai_settings.reset_defaults",
            )

        ctk.CTkButton(btns, text="Save", command=save).pack(side="right", padx=6)
        ctk.CTkButton(btns, text="Defaults", command=reset_defaults).pack(side="right", padx=6)
        ctk.CTkButton(btns, text="Close", command=top.destroy).pack(side="right", padx=6)

    # ---------------------------
    # Setup and Layout Methods
    # ---------------------------
    def set_window_icon(self):
        icon_path = os.path.join("assets", "GMCampaignDesigner.ico")
        if os.path.exists(icon_path):
            self.iconbitmap(icon_path)
        icon_image = PhotoImage(file=os.path.join("assets", "GMCampaignDesigner logo.png"))
        self.tk.call('wm', 'iconphoto', self._w, icon_image)

    def create_layout(self):
        self.main_frame = ctk.CTkFrame(self)
        self.main_frame.pack(fill="both", expand=True)

    def load_icons(self):
        log_debug("Loading sidebar icons", func_name="main_window.MainWindow.load_icons")
        base_icons = {
            "change_db": "database_icon.png",
            "db_export": "database_export_icon.png",
            "db_import": "database_import_icon.png",
            "swarm_path": "folder_icon.png",
            "customize_fields": "customize_fields.png",
            "new_entity_type": "customize_fields.png",
            "export_scenarios": "export_icon.png",
            "asset_library": "icons/save.png",
            "gm_screen": "gm_screen_icon.png",
            "npc_graph": "npc_graph_icon.png",
            "pc_graph": "pc_graph_icon.png",
            "faction_graph": "faction_graph_icon.png",
            "scenario_graph": "scenario_graph_icon.png",
            "scenario_builder": "scenario_graph_icon.png",
            "world_map": "maps_icon.png",
            "generate_portraits": "generate_icon.png",
            "associate_portraits": "associate_icon.png",
            "import_portraits": "import_icon.png",
            "import_scenario": "import_icon.png",
            "import_creatures_pdf": "import_icon.png",
            "import_objects_pdf": "import_icon.png",
            "export_foundry": "export_foundry_icon.png",
            "map_tool": "map_tool_icon.png",
            "whiteboard": "map_tool_icon.png",
            "generate_scenario": "generate_scenario_icon.png",
            "dice_roller": "dice_roller_icon.png",
            "dice_bar": "dice_roller_icon.png",
            "sound_manager": "sound_manager_icon.png",
            "audio_controls": "sound_manager_icon.png",
            "scene_flow_viewer": "scenes_flow_icon.png",
            "create_random_table"   : "random_table_icon.png",
        }

        # Entity-specific icons come from metadata; keep base ones separate.
        self.icons = {}
        for key, file_name in base_icons.items():
            self.icons[key] = self.load_icon(file_name, size=(60, 60))

        default_entity_icon = self.icons.get("customize_fields") or self.load_icon("customize_fields.png", size=(60, 60))
        for slug, meta in self.entity_definitions.items():
            icon_key = f"entity::{slug}"
            icon_path = meta.get("icon")
            icon_image = self.load_icon(icon_path, size=(60, 60)) if icon_path else None
            self.icons[icon_key] = icon_image or default_entity_icon

    def open_custom_fields_editor(self):
        try:
            log_info("Opening Custom Fields Editor", func_name="main_window.MainWindow.open_custom_fields_editor")
            top = CustomFieldsEditor(self)
            top.transient(self)
            top.lift(); top.focus_force()
        except Exception as e:
            log_exception(f"Failed to open Custom Fields Editor: {e}",
                          func_name="main_window.MainWindow.open_custom_fields_editor")
            messagebox.showerror("Error", f"Failed to open Custom Fields Editor:\n{e}")

    def open_cross_campaign_asset_library(self):
        try:
            if self._asset_library_window and self._asset_library_window.winfo_exists():
                self._asset_library_window.lift()
                self._asset_library_window.focus_force()
                return
            log_info(
                "Opening Cross-campaign Asset Library",
                func_name="main_window.MainWindow.open_cross_campaign_asset_library",
            )
            window = CrossCampaignAssetLibraryWindow(self)
            window.bind("<Destroy>", lambda _evt: setattr(self, "_asset_library_window", None))
            self._asset_library_window = window
        except Exception as exc:
            log_exception(
                f"Failed to open asset library: {exc}",
                func_name="main_window.MainWindow.open_cross_campaign_asset_library",
            )
            messagebox.showerror("Error", f"Failed to open asset library window:\n{exc}")

    def open_new_entity_type_dialog(self):
        try:
            log_info("Opening New Entity Type dialog", func_name="main_window.MainWindow.open_new_entity_type_dialog")
            dlg = NewEntityTypeDialog(self, on_created=self.refresh_entities)
            dlg.transient(self)
            dlg.lift(); dlg.focus_force()
        except Exception as exc:
            log_exception(
                f"Failed to open New Entity Type dialog: {exc}",
                func_name="main_window.MainWindow.open_new_entity_type_dialog",
            )
            messagebox.showerror("Error", f"Failed to open New Entity Type dialog:\n{exc}")

    def load_icon(self, file_name, size=(60, 60)):
        if not file_name:
            return None

        candidates = []
        if os.path.isabs(file_name):
            candidates.append(file_name)
        else:
            if os.path.exists(file_name):
                candidates.append(file_name)
            candidates.append(os.path.join("assets", file_name))
            campaign_dir = ConfigHelper.get_campaign_dir()
            candidates.append(os.path.join(campaign_dir, file_name))
            candidates.append(os.path.join(campaign_dir, "assets", file_name))
            if os.path.sep in file_name:
                candidates.append(file_name)

        seen = set()
        for path in candidates:
            if not path or path in seen:
                continue
            seen.add(path)
            if not os.path.exists(path):
                continue
            try:
                pil_image = Image.open(path)
            except Exception as exc:
                log_warning(
                    f"Unable to load icon {path}: {exc}",
                    func_name="main_window.MainWindow.load_icon",
                )
                continue
            log_debug(f"Loaded icon {path}", func_name="main_window.MainWindow.load_icon")
            return ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=size)

        log_warning(
            f"Unable to locate icon {file_name}",
            func_name="main_window.MainWindow.load_icon",
        )
        return None

    def create_sidebar(self, force: bool = False):
        first_time = not hasattr(self, "sidebar_frame") or force

        if first_time:
            self._ensure_sidebar_hotspot()
            # If forcing a rebuild, drop the existing sidebar completely
            if force and hasattr(self, "sidebar_frame"):
                try:
                    self.sidebar_frame.destroy()
                except Exception:
                    pass
            self.sidebar_frame = ctk.CTkFrame(self.main_frame, width=self.sidebar_default_width)
            pack_kwargs = {"side": "left", "fill": "y", "padx": 5, "pady": 5}
            if force and hasattr(self, "content_frame"):
                try:
                    if self.content_frame.winfo_manager() == "pack":
                        pack_kwargs["before"] = self.content_frame
                except Exception:
                    pass
            if getattr(self, "sidebar_hotspot", None) is not None:
                pack_kwargs.setdefault("after", self.sidebar_hotspot)
            self.sidebar_frame.pack(**pack_kwargs)
            self.sidebar_frame.pack_propagate(False)
            self._sidebar_pack_kwargs = pack_kwargs
            self.sidebar_inner = ctk.CTkFrame(self.sidebar_frame, fg_color="transparent")
            self.sidebar_inner.pack(fill="both", expand=True, padx=5, pady=5)

            # Ensure the banner toggle button persists if the sidebar is rebuilt
            if hasattr(self, "banner_toggle_btn"):
                existing_btn = self.banner_toggle_btn

                try:
                    exists = bool(existing_btn.winfo_exists())
                except Exception:
                    exists = False

                # If the previous widget survived, simply re-place it; otherwise rebuild it
                if not exists or existing_btn.master is not self.sidebar_inner:
                    if exists:
                        try:
                            existing_btn.destroy()
                        except Exception:
                            pass
                    self.banner_toggle_btn = ctk.CTkButton(
                        self.sidebar_inner,
                        text="▼",
                        width=40,
                        height=30,
                        command=self._toggle_banner,
                        fg_color="#555",
                        hover_color="#777",
                        font=("", 16),
                    )
                    self.banner_toggle_btn.place(relx=1.0, rely=0.0, anchor="ne")
                else:
                    try:
                        manager_name = existing_btn.winfo_manager()
                    except Exception:
                        manager_name = ""
                    if not manager_name:
                        log_debug(
                            "Banner toggle button lost geometry manager; reattaching.",
                            func_name="main_window.MainWindow.create_sidebar",
                        )
                    # Calling place ensures the toggle reattaches to the refreshed sidebar
                    try:
                        existing_btn.place(relx=1.0, rely=0.0, anchor="ne")
                    except Exception:
                        pass

            # Logo
            logo_path = os.path.join("assets", "GMCampaignDesigner logo.png")
            if os.path.exists(logo_path):
                logo_image = Image.open(logo_path).resize((60, 60))
                logo = ctk.CTkImage(light_image=logo_image, dark_image=logo_image, size=(80, 80))
                self.logo_image = logo
                logo_label = ctk.CTkLabel(self.sidebar_inner, image=logo, text="")
                logo_label.pack(pady=(0, 3), anchor="center")

            # Header label
            header_label = ctk.CTkLabel(self.sidebar_inner, text="Campaign Tools", font=("Helvetica", 16, "bold"))
            header_label.pack(pady=(0, 2), anchor="center")

            # Database display container
            tokens = theme_manager.get_tokens()
            db_container = ctk.CTkFrame(
                self.sidebar_inner,
                fg_color="transparent",
                border_color=tokens.get("button_border"),
                border_width=2,
                corner_radius=8,
            )
            db_container.pack(pady=(0, 5), anchor="center", fill="x", padx=5)
            db_title_label = ctk.CTkLabel(
                db_container,
                text="Database:",
                font=("Segoe UI", 16, "bold"),
                fg_color="transparent",
                text_color="white",
            )
            db_title_label.pack(pady=(3, 0), anchor="center")
            db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
            db_name = os.path.splitext(os.path.basename(db_path))[0]
            self.db_name_label = ctk.CTkLabel(
                db_container,
                text=db_name,
                font=("Segoe UI", 14, "italic"),
                fg_color="transparent",
                text_color="white",
            )
            self.db_name_label.pack(pady=(0, 1), anchor="center")
            try:
                full_path = os.path.abspath(db_path)
                self.db_tooltip = ToolTip(self.db_name_label, full_path)
            except Exception:
                self.db_tooltip = None

            self.system_label = ctk.CTkLabel(
                db_container,
                text="System: --",
                font=("Segoe UI", 13),
                fg_color="transparent",
                text_color="#d5e6ff",
            )
            self.system_label.pack(pady=(0, 2), anchor="center")
            self.system_tooltip = None

            sys_tokens = theme_manager.get_tokens()
            system_button = ctk.CTkButton(
                db_container,
                text="Switch System",
                command=self.open_system_selector,
                width=160,
                fg_color=sys_tokens.get("accent_button_fg"),
                hover_color=sys_tokens.get("accent_button_hover"),
            )
            system_button.pack(pady=(0, 6))
            self.system_button = system_button

            self.sidebar_sections_container = ctk.CTkFrame(self.sidebar_inner, fg_color="transparent")
            self.sidebar_sections_container.pack(fill="both", expand=True, padx=5, pady=5)
        else:
            # Refresh top metadata styles to pick up new theme tokens
            try:
                tokens = theme_manager.get_tokens()
                # Update border color on DB container if present
                for child in self.sidebar_inner.winfo_children():
                    try:
                        if isinstance(child, ctk.CTkFrame) and getattr(child, "border_width", 0):
                            child.configure(border_color=tokens.get("button_border"))
                    except Exception:
                        pass
                if getattr(self, "system_button", None) is not None:
                    try:
                        self.system_button.configure(
                            fg_color=tokens.get("accent_button_fg"),
                            hover_color=tokens.get("accent_button_hover"),
                        )
                    except Exception:
                        pass
            except Exception:
                pass
            for child in self.sidebar_sections_container.winfo_children():
                child.destroy()

        self._bind_sidebar_hover_events()
        self.update_sidebar_metadata()

        self.create_accordion_sidebar()

        if self._sidebar_collapsed:
            self._collapse_sidebar(immediate=True)


    def update_sidebar_metadata(self):
        db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
        db_name = os.path.splitext(os.path.basename(db_path))[0]

        label = getattr(self, "db_name_label", None)
        if label is not None and label.winfo_exists():
            label.configure(text=db_name)

        tooltip = getattr(self, "db_tooltip", None)
        if tooltip is not None:
            try:
                tooltip.text = os.path.abspath(db_path)
            except Exception:
                tooltip.text = db_path

        config = system_config.get_current_system_config()
        system_label_text = "Unknown"
        system_slug = None
        if config is not None:
            system_label_text = config.label or config.slug or "Unknown"
            system_slug = config.slug

        system_label = getattr(self, "system_label", None)
        if system_label is not None and system_label.winfo_exists():
            system_label.configure(text=f"System: {system_label_text}")

        system_tooltip = getattr(self, "system_tooltip", None)
        tooltip_text = (
            f"Active slug: {system_slug}" if system_slug else "No system selected"
        )
        if system_tooltip is None and system_label is not None:
            try:
                self.system_tooltip = ToolTip(system_label, tooltip_text)
            except Exception:
                self.system_tooltip = None
        elif system_tooltip is not None:
            system_tooltip.text = tooltip_text

    def open_system_selector(self):
        existing = getattr(self, "_system_selector_dialog", None)
        if existing is not None and existing.winfo_exists():
            try:
                existing.lift()
                existing.focus_force()
            except Exception:
                pass
            return

        dialog = CampaignSystemSelectorDialog(self)
        dialog.bind("<Destroy>", self._on_system_selector_destroyed)
        self._system_selector_dialog = dialog

    def _on_system_selector_destroyed(self, event=None):
        if event is None:
            self._system_selector_dialog = None
            return
        if event.widget is self._system_selector_dialog:
            self._system_selector_dialog = None

    def _on_theme_changed(self, _theme_key: str) -> None:
        # Reapply palette and rebuild widgets that used tokenized colors
        # Track which floating windows are currently open so we can recreate them
        reopen_audio_bar = False
        reopen_dice_bar = False
        reopen_dice_roller = False
        reopen_sound_manager = False
        try:
            if getattr(self, "audio_bar_window", None) is not None and self.audio_bar_window.winfo_exists():
                reopen_audio_bar = True
                try:
                    self.audio_bar_window.destroy()
                except Exception:
                    pass
        except Exception:
            pass
        try:
            if getattr(self, "dice_bar_window", None) is not None and self.dice_bar_window.winfo_exists():
                reopen_dice_bar = True
                try:
                    self.dice_bar_window.destroy()
                except Exception:
                    pass
        except Exception:
            pass
        try:
            if getattr(self, "dice_roller_window", None) is not None and self.dice_roller_window.winfo_exists():
                reopen_dice_roller = True
                try:
                    self.dice_roller_window.destroy()
                except Exception:
                    pass
        except Exception:
            pass
        try:
            if getattr(self, "sound_manager_window", None) is not None and self.sound_manager_window.winfo_exists():
                reopen_sound_manager = True
                try:
                    self.sound_manager_window.destroy()
                except Exception:
                    pass
        except Exception:
            pass

        try:
            theme_manager.apply_theme(_theme_key)
        except Exception:
            pass
        try:
            # Recreate sidebar so header bands and borders update
            self.create_sidebar(force=True)
        except Exception:
            # Fallback: at least rebuild the sections area
            try:
                self.create_accordion_sidebar()
            except Exception:
                pass

        # Re-open floating bars/windows to pick up the new palette immediately
        try:
            if reopen_audio_bar:
                self.after(25, self.open_audio_bar)
        except Exception:
            pass
        try:
            if reopen_dice_bar:
                self.after(25, self.open_dice_bar)
        except Exception:
            pass
        try:
            if reopen_dice_roller:
                self.after(50, self.open_dice_roller)
        except Exception:
            pass
        try:
            if reopen_sound_manager:
                self.after(75, self.open_sound_manager)
        except Exception:
            pass

    def create_accordion_sidebar(self):
        container = getattr(self, "sidebar_sections_container", None)
        if container is None:
            return
        for child in container.winfo_children():
            child.destroy()
        sections = []  # track all sections to enforce single-open behavior
        default_title = "Campaign Workshop"
        default_meta = {"sec": None}
        active_section = {"sec": None}

        def make_section(parent, title, buttons):
            sec = ctk.CTkFrame(parent)
            sec.pack(fill="x", pady=(4, 6))

            header_color = theme_manager.get_tokens().get("sidebar_header_bg")
            header = ctk.CTkFrame(sec, fg_color=header_color)
            header.pack(fill="x")
            title_lbl = ctk.CTkLabel(header, text=title, anchor="center")
            title_lbl.pack(fill="x", pady=6)
            state = {"open": False}

            body = ctk.CTkFrame(sec, fg_color="transparent")
            cols = 2
            for c in range(cols):
                body.grid_columnconfigure(c, weight=1)
            for idx, (icon_key, tooltip, cmd) in enumerate(buttons):
                r, c = divmod(idx, cols)
                icon = self.icons.get(icon_key)
                btn = create_icon_button(body, icon, tooltip, cmd)
                btn.grid(row=r, column=c, padx=2, pady=2, sticky="ew")

            def expand():
                if state["open"]:
                    return
                state["open"] = True
                try:
                    body.pack(fill="x", padx=2, pady=(2, 2))
                except Exception:
                    pass

            def collapse():
                if not state["open"]:
                    return
                state["open"] = False
                try:
                    body.pack_forget()
                except Exception:
                    pass

            def toggle(_event=None):
                if state["open"]:
                    collapse()
                    if active_section["sec"] is sec:
                        active_section["sec"] = None
                else:
                    for meta in sections:
                        if meta["sec"] is sec:
                            continue
                        meta["collapse"]()
                    expand()
                    active_section["sec"] = sec

            header.bind("<Button-1>", toggle)
            title_lbl.bind("<Button-1>", toggle)

            sections.append({
                "sec": sec,
                "state": state,
                "collapse": collapse,
                "expand": expand,
                "toggle": toggle,
                "title": title,
            })
            return sec

        # Group buttons
        data_system = [
            ("change_db", "Change Data Storage", self.change_database_storage),
            ("swarm_path", "Set SwarmUI Path", self.select_swarmui_path),
            ("customize_fields", "Customize Fields", self.open_custom_fields_editor),
            ("new_entity_type", "New Entity Type", self.open_new_entity_type_dialog),
            ("db_export", "Create Campaign Backup", self.prompt_campaign_backup),
            ("db_import", "Restore Campaign Backup", self.prompt_campaign_restore),
            ("asset_library", "Cross-campaign Asset Library", self.open_cross_campaign_asset_library),
        ]

        entity_buttons = []
        for slug, meta in self.entity_definitions.items():
            key = f"entity::{slug}"
            label = meta.get("label") or slug.replace("_", " ").title()
            tooltip = f"Manage {label}"
            entity_buttons.append((key, tooltip, lambda s=slug: self.open_entity(s)))

        relations = [
            ("npc_graph", "Open NPC Graph Editor", self.open_npc_graph_editor),
            ("pc_graph", "Open PC Graph Editor", self.open_pc_graph_editor),
            ("faction_graph", "Open Factions Graph Editor", self.open_faction_graph_editor),
            ("scenario_graph", "Open Scenario Graph Editor", self.open_scenario_graph_editor),
            ("create_random_table", "Create Random Table", self.open_random_table_editor),
            ("scene_flow_viewer", "Open Scene Flow Viewer", self.open_scene_flow_viewer),
            ("world_map", "Open World Map", self.open_world_map),
        ]
        utilities = [
            ("generate_scenario", "Generate Scenario", self.open_scenario_generator),
            ("scenario_builder", "Scenario Builder Wizard", self.open_scenario_builder),
            ("import_scenario", "Import Scenario", self.open_scenario_importer),
            ("import_creatures_pdf", "Import Creatures from PDF", self.open_creature_importer),
            ("import_objects_pdf", "Import Equipment from PDF", self.open_object_importer),
            ("gm_screen", "Open GM Screen", self.open_gm_screen),
            ("export_scenarios", "Export Scenarios", self.preview_and_export_scenarios),
            ("export_foundry", "Export Scenarios for Foundry", self.export_foundry),
            ("generate_portraits", "Generate Portraits", self.generate_missing_portraits),
            ("associate_portraits", "Associate NPC Portraits", self.associate_npc_portraits),
            ("import_portraits", "Import Portraits from Folder", self.import_portraits_from_directory),
            ("map_tool", "Map Tool", self.map_tool),
            ("whiteboard", "Whiteboard", self.open_whiteboard),
            ("sound_manager", "Sound & Music Manager", self.open_sound_manager),
            ("dice_roller", "Open Dice Roller", self.open_dice_roller),
        ]

        make_section(container, "Data & System", data_system)
        make_section(container, "Campaign Workshop", entity_buttons)
        make_section(container, "Relations & Graphs", relations)
        make_section(container, "Utilities", utilities)

        # Open the default section by default
        for meta in sections:
            if meta.get("title") == default_title:
                default_meta = meta
                try:
                    meta["expand"]()
                    active_section["sec"] = meta["sec"]
                except Exception:
                    pass
                break

    def _ensure_sidebar_hotspot(self):
        if getattr(self, "sidebar_hotspot", None) is not None:
            return
        self.sidebar_hotspot = ctk.CTkFrame(self.main_frame, width=8, fg_color="transparent")
        self.sidebar_hotspot.pack(side="left", fill="y")
        try:
            self.sidebar_hotspot.bind("<Enter>", self._restore_sidebar)
        except Exception:
            pass

    def _bind_sidebar_hover_events(self):
        frame = getattr(self, "sidebar_frame", None)
        if frame is None:
            return
        try:
            frame.bind("<Enter>", self._restore_sidebar)
            frame.bind("<Leave>", self._collapse_sidebar)
            if getattr(self, "sidebar_inner", None) is not None:
                self.sidebar_inner.bind("<Enter>", self._restore_sidebar)
                self.sidebar_inner.bind("<Leave>", self._collapse_sidebar)
        except Exception:
            pass

    def _cancel_sidebar_animation(self):
        if self._sidebar_animation_job is not None:
            try:
                self.after_cancel(self._sidebar_animation_job)
            except Exception:
                pass
        self._sidebar_animation_job = None
        self._sidebar_animating = False

    def _animate_sidebar_width(self, target_width: int, on_complete=None):
        frame = getattr(self, "sidebar_frame", None)
        if frame is None or not frame.winfo_exists():
            self._sidebar_animating = False
            return

        # Cancel any pending animation callbacks before starting a fresh one
        self._cancel_sidebar_animation()

        # Immediately set the target width to remove the animation effect
        try:
            frame.configure(width=target_width)
        except Exception:
            self._sidebar_animating = False
            if callable(on_complete):
                on_complete()
            return

        self._sidebar_animating = False
        if callable(on_complete):
            on_complete()

    def _restore_sidebar(self, _event=None):
        frame = getattr(self, "sidebar_frame", None)
        if frame is None or not frame.winfo_exists():
            return
        self._cancel_sidebar_animation()

        if not frame.winfo_manager():
            if self._sidebar_pack_kwargs is None:
                self._sidebar_pack_kwargs = {"side": "left", "fill": "y", "padx": 5, "pady": 5}
                if getattr(self, "sidebar_hotspot", None) is not None:
                    self._sidebar_pack_kwargs["after"] = self.sidebar_hotspot
            try:
                frame.configure(width=0)
                frame.pack(**self._sidebar_pack_kwargs)
                frame.pack_propagate(False)
            except Exception:
                return

        target_width = self.sidebar_default_width

        def finalize():
            self._sidebar_collapsed = False
            try:
                frame.configure(width=target_width)
            except Exception:
                pass

        self._animate_sidebar_width(target_width, on_complete=finalize)

    def _is_pointer_exiting_sidebar_right(self, event) -> bool:
        """Return True only when the pointer exits to the right of the sidebar frame."""
        frame = getattr(self, "sidebar_frame", None)
        if frame is None or not frame.winfo_exists():
            return True

        # Allow explicit collapses (no event) to proceed
        if event is None:
            return True

        try:
            frame_left = frame.winfo_rootx()
            frame_right = frame_left + frame.winfo_width()
        except Exception:
            return True

        root_x = getattr(event, "x_root", None)
        if root_x is None:
            return False

        if root_x >= frame_right:
            return True

        # Pointer is still over or to the left of the sidebar; do not collapse yet
        return False

    def _collapse_sidebar(self, _event=None, immediate: bool = False):
        frame = getattr(self, "sidebar_frame", None)
        if frame is None or not frame.winfo_exists():
            return
        self._cancel_sidebar_animation()
        if self._sidebar_collapsed:
            return

        if _event is not None and not immediate:
            if not self._is_pointer_exiting_sidebar_right(_event):
                return

        def finalize():
            try:
                frame.pack_forget()
            except Exception:
                pass
            self._sidebar_collapsed = True
            self._sidebar_animating = False

        if immediate:
            finalize()
            return

        self._animate_sidebar_width(0, on_complete=finalize)

    def create_content_area(self):
        self.content_frame = ctk.CTkFrame(self.main_frame)
        self.content_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)

        self.content_frame.grid_rowconfigure(0, weight=1)
        self.content_frame.grid_rowconfigure(1, weight=0)
        self.content_frame.grid_columnconfigure(0, weight=1)

        # ✅ Always explicitly create these initially:
        self.banner_frame = ctk.CTkFrame(self.content_frame, height=150, fg_color="#444")
        # — Monkey-patch banner_frame.pack to grid instead, so open_gm_screen’s pack() won’t clash
        self.banner_frame.pack = lambda *args, **kwargs: self.banner_frame.grid(row=0, column=0, sticky="ew")
        self.inner_content_frame = ctk.CTkFrame(self.content_frame, fg_color="#222")

        self.banner_toggle_btn = ctk.CTkButton(
            self.sidebar_inner,
            text="▼",
            width=40,
            height=30,
            command=self._toggle_banner,
            fg_color="#555",
            hover_color="#777",
            font=("", 16)
        )
        self.banner_toggle_btn.place(relx=1.0, rely=0.0, anchor="ne")

        self.banner_visible = False
        self.current_open_view = None

    def _prime_content_frames_for_gm_screen(self):
        """Normalize banner/content grid so GM Screen layout starts in a known state."""
        try:
            self.content_frame.grid_rowconfigure(0, weight=0)
            self.content_frame.grid_rowconfigure(1, weight=1)
            self.content_frame.grid_columnconfigure(0, weight=1)
        except Exception:
            pass

        try:
            self.banner_frame.grid(row=0, column=0, sticky="ew")
        except Exception:
            pass

        try:
            self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
        except Exception:
            pass

        self.banner_visible = True
        try:
            self.banner_toggle_btn.configure(text="▲")
        except Exception:
            pass

    def _toggle_banner(self):
        # GM Screen mode: reposition content and toggle banner without recreating views
        if getattr(self, "_gm_mode", False):
            if self.banner_visible:
                if self.banner_frame.winfo_exists():
                    try:
                        self.banner_frame.grid_remove()
                    except Exception:
                        pass
                try:
                    self.inner_content_frame.grid(row=0, column=0, sticky="nsew")
                except Exception:
                    pass
                self.content_frame.grid_rowconfigure(0, weight=1)
                self.content_frame.grid_rowconfigure(1, weight=0)
                self.banner_visible = False
                try:
                    self.banner_toggle_btn.configure(text="▼")
                except Exception:
                    pass
            else:
                try:
                    self.banner_frame.grid(row=0, column=0, sticky="ew")
                except Exception:
                    pass
                try:
                    self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
                except Exception:
                    pass
                try:
                    display_pcs_in_banner(
                        self.banner_frame,
                        {pc["Name"]: pc for pc in self.pc_wrapper.load_items()}
                    )
                except Exception:
                    pass
                self.inner_content_frame.grid_rowconfigure(0, weight=1)
                self.inner_content_frame.grid_columnconfigure(0, weight=1)
                self.content_frame.grid_rowconfigure(0, weight=0)
                self.content_frame.grid_rowconfigure(1, weight=1)
                self.banner_visible = True
                try:
                    self.banner_toggle_btn.configure(text="▲")
                except Exception:
                    pass
            return
        # --- GRAPH MODE?  (no current_open_entity but a graph_type set) ---
        if self.current_open_entity is None and getattr(self, "_graph_type", None):
            # snapshot existing graph state
            old_container = self.current_open_view
            old_editor    = getattr(old_container, "graph_editor", None)
            state         = old_editor.get_state() if old_editor else None

            # hide or show banner + inner frames
            if self.banner_visible:
                self.banner_frame.grid_remove()
                self.inner_content_frame.grid_remove()
                self.content_frame.grid_rowconfigure(0, weight=1)
                self.content_frame.grid_rowconfigure(1, weight=0)
                self.banner_visible = False
                self.banner_toggle_btn.configure(text="▼")
            else:
                self.banner_frame.grid(row=0, column=0, sticky="ew")
                self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
                display_pcs_in_banner(
                    self.banner_frame,
                    {pc["Name"]: pc for pc in self.pc_wrapper.load_items()}
                )
                self.inner_content_frame.grid_rowconfigure(0, weight=1)
                self.inner_content_frame.grid_columnconfigure(0, weight=1)
                self.content_frame.grid_rowconfigure(0, weight=0)
                self.content_frame.grid_rowconfigure(1, weight=1)
                self.banner_visible = True
                self.banner_toggle_btn.configure(text="▲")

            # destroy the old container (it was still parented in the wrong frame)
            old_container.destroy()

            # now re‐create the same graph under the correct container
            parent = self.get_content_container()
            new_container = ctk.CTkFrame(parent)
            new_container.grid(row=0, column=0, sticky="nsew")
            parent.grid_rowconfigure(0, weight=1)
            parent.grid_columnconfigure(0, weight=1)

            # re‐instantiate the proper editor type, then restore its state
            if self._graph_type == 'npc':
                editor = NPCGraphEditor(new_container, self.npc_wrapper, self.faction_wrapper)
            elif self._graph_type == 'pc':
                editor = PCGraphEditor(new_container, self.pc_wrapper, self.faction_wrapper)
            elif self._graph_type == 'faction':
                editor = FactionGraphEditor(new_container, self.faction_wrapper)
            else:  # 'scenario'
                editor = ScenarioGraphEditor(
                    new_container,
                    GenericModelWrapper("scenarios"),
                    GenericModelWrapper("npcs"),
                    GenericModelWrapper("creatures"),
                    GenericModelWrapper("places")
                )

            editor.pack(fill="both", expand=True)
            if state is not None and hasattr(editor, "set_state"):
                editor.set_state(state)

            # save the new container/editor
            new_container.graph_editor = editor
            self.current_open_view     = new_container
            # leave current_open_entity = None

            return  # end of graph‐mode toggle
        if self.banner_visible:
            # COLLAPSE BANNER
            if self.banner_frame.winfo_exists():
                self.banner_frame.grid_remove()
            if self.inner_content_frame.winfo_exists():
                self.inner_content_frame.grid_remove()

            if self.current_open_view:
                # Save entity and then destroy
                entity = self.current_open_entity
                self.current_open_view.destroy()

                # Re-create content in content_frame
                self.current_open_view = ctk.CTkFrame(self.content_frame)
                self.current_open_view.grid(row=0, column=0, sticky="nsew")

                wrapper = self.entity_wrappers.get(entity)
                if wrapper is None:
                    wrapper = GenericModelWrapper(entity)
                    self.entity_wrappers[entity] = wrapper
                template = load_template(entity)
                view = GenericListView(self.current_open_view, wrapper, template)
                view.pack(fill="both", expand=True)

                meta = self.entity_definitions.get(entity, {})
                display_label = meta.get("label") or entity.replace("_", " ").title()

                load_button = ctk.CTkButton(
                    self.current_open_view,
                    text=f"Load {display_label}",
                    command=lambda: self.load_items_from_json(view, entity)
                )
                load_button.pack(side="right", padx=(5,5))
                # Assuming `editor_window` is your CTkToplevel or CTkFrame
                save_btn = ctk.CTkButton(
                    self.current_open_view,
                    text=f"Save {display_label}",
                    command=lambda: self.save_items_to_json(view, entity)
                )
                save_btn.pack(side="right", padx=(5,5))

            self.content_frame.grid_rowconfigure(0, weight=1)
            self.content_frame.grid_rowconfigure(1, weight=0)

            self.banner_visible = False
            self.banner_toggle_btn.configure(text="▼")
        else:
            # EXPAND BANNER
            if not self.banner_frame.winfo_exists():
                self.banner_frame = ctk.CTkFrame(self.content_frame, height=150, fg_color="#444")

            if not self.inner_content_frame.winfo_exists():
                self.inner_content_frame = ctk.CTkFrame(self.content_frame, fg_color="#222")

            self.banner_frame.grid(row=0, column=0, sticky="ew")
            self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
            pcs_items = {pc["Name"]: pc for pc in self.pc_wrapper.load_items()}
            display_pcs_in_banner(self.banner_frame, pcs_items)

            # ✅ CRITICAL FIX: make inner_content_frame fully expandable
            self.inner_content_frame.grid_rowconfigure(0, weight=1)
            self.inner_content_frame.grid_columnconfigure(0, weight=1)

            if self.current_open_view:
                entity = self.current_open_entity
                self.current_open_view.destroy()

                self.current_open_view = ctk.CTkFrame(self.inner_content_frame)
                self.current_open_view.grid(row=0, column=0, sticky="nsew")

                wrapper = self.entity_wrappers.get(entity)
                if wrapper is None:
                    wrapper = GenericModelWrapper(entity)
                    self.entity_wrappers[entity] = wrapper
                template = load_template(entity)
                view = GenericListView(self.current_open_view, wrapper, template)
                view.pack(fill="both", expand=True)

                meta = self.entity_definitions.get(entity, {})
                display_label = meta.get("label") or entity.replace("_", " ").title()

                load_button = ctk.CTkButton(
                    self.current_open_view,
                    text=f"Load {display_label}",
                    command=lambda: self.load_items_from_json(view, entity)
                )
                load_button.pack(side="right", padx=(5,5))
                # Assuming `editor_window` is your CTkToplevel or CTkFrame
                save_btn = ctk.CTkButton(
                    self.current_open_view,
                    text=f"Save {display_label}",
                    command=lambda: self.save_items_to_json(view, entity)
                )
                save_btn.pack(side="right", padx=(5,5))

            self.content_frame.grid_rowconfigure(0, weight=0)
            self.content_frame.grid_rowconfigure(1, weight=1)

            self.banner_visible = True
            self.banner_toggle_btn.configure(text="▲")

    def get_content_container(self):
        """Choose correct parent depending on banner state."""
        if self.banner_visible:
            return self.inner_content_frame
        else:
            return self.content_frame
    def create_exit_button(self):
        exit_button = ctk.CTkButton(self, text="✕", command=self.destroy,
                                    fg_color="red", hover_color="#AA0000",
                                    width=20, height=20, corner_radius=15)
        exit_button.place(relx=0.9999, rely=0.01, anchor="ne")

    def _apply_cursor_recursive(self, widget, cursor):
        try:
            widget.configure(cursor=cursor)
        except (tk.TclError, AttributeError):
            pass
        for child in widget.winfo_children():
            self._apply_cursor_recursive(child, cursor)

    def _set_wait_cursor(self, enable):
        cursor = "wait" if sys.platform.startswith("win") else "watch"
        target = cursor if enable else ""
        try:
            self._apply_cursor_recursive(self, target)
        except tk.TclError:
            pass
        if self._busy_modal and self._busy_modal.winfo_exists():
            try:
                self._busy_modal.configure(cursor=target)
            except tk.TclError:
                pass
        self.update_idletasks()

    def _show_busy_modal(self, message):
        if self._busy_modal and self._busy_modal.winfo_exists():
            try:
                label = getattr(self._busy_modal, "_message_label", None)
                if label is not None:
                    label.configure(text=message)
                self._busy_modal.lift()
                return
            except tk.TclError:
                self._busy_modal = None

        modal = ctk.CTkToplevel(self)
        modal.title("Please wait")
        modal.resizable(False, False)
        modal.transient(self)
        modal.protocol("WM_DELETE_WINDOW", lambda: None)

        frame = ctk.CTkFrame(modal)
        frame.pack(fill="both", expand=True, padx=20, pady=20)

        label = ctk.CTkLabel(frame, text=message, anchor="center")
        label.pack(fill="both", expand=True)
        modal._message_label = label  # type: ignore[attr-defined]

        modal.update_idletasks()
        self.update_idletasks()
        try:
            width = modal.winfo_width()
            height = modal.winfo_height()
            root_x = self.winfo_rootx()
            root_y = self.winfo_rooty()
            root_w = self.winfo_width()
            root_h = self.winfo_height()
            x = root_x + max((root_w - width) // 2, 0)
            y = root_y + max((root_h - height) // 2, 0)
            modal.geometry(f"{width}x{height}+{x}+{y}")
        except tk.TclError:
            pass

        modal.attributes("-topmost", True)
        modal.after(10, lambda: modal.attributes("-topmost", False))
        self._busy_modal = modal
        self._set_wait_cursor(True)

    def _hide_busy_modal(self):
        try:
            if self._busy_modal and self._busy_modal.winfo_exists():
                self._busy_modal.destroy()
        except tk.TclError:
            pass
        finally:
            self._busy_modal = None
            self._set_wait_cursor(False)

    @contextmanager
    def _busy_operation(self, message):
        self._show_busy_modal(message)
        try:
            yield
        finally:
            self._hide_busy_modal()

    def load_model_config(self):
        self.models_path = ConfigHelper.get("Paths", "models_path",
                                            fallback=r"E:\SwarmUI\SwarmUI\Models\Stable-diffusion")
        self.model_options = get_available_models()
        log_debug(
            f"Models path resolved to {self.models_path}",
            func_name="main_window.MainWindow.load_model_config",
        )
        log_info(
            f"Loaded {len(self.model_options)} available AI models",
            func_name="main_window.MainWindow.load_model_config",
        )

    def init_wrappers(self):
        self.entity_wrappers = build_entity_wrappers()

        attr_map = {
            "places": ["place_wrapper"],
            "npcs": ["npc_wrapper"],
            "pcs": ["pc_wrapper"],
            "factions": ["faction_wrapper"],
            "objects": ["object_wrapper"],
            "creatures": ["creature_wrapper"],
            "clues": ["clues_wrapper"],
            "books": ["books_wrapper"],
            "informations": ["informations_wrapper", "information_wrapper"],
            "maps": ["maps_wrapper"],
        }

        for slug, attrs in attr_map.items():
            wrapper = self.entity_wrappers.get(slug)
            if wrapper is None:
                wrapper = GenericModelWrapper(slug)
                self.entity_wrappers[slug] = wrapper
            for attr in attrs:
                setattr(self, attr, wrapper)

        log_info("Entity wrappers initialized", func_name="main_window.MainWindow.init_wrappers")

    def _normalize_entity_media_paths(self):
        """Ensure stored portrait/image paths are relative within the campaign."""

        raw_db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db") or "default_campaign.db"
        campaign_dir_candidates = []

        try:
            campaign_dir_candidates.append(Path(raw_db_path).parent)
        except Exception:
            pass

        try:
            campaign_dir_candidates.append(Path(ConfigHelper.get_campaign_dir()))
        except Exception:
            pass

        # Deduplicate while preserving order
        unique_candidates = []
        seen = set()
        for candidate in campaign_dir_candidates:
            candidate_str = str(candidate)
            if candidate_str and candidate_str not in seen:
                unique_candidates.append(candidate)
                seen.add(candidate_str)

        path_fields = ("Portrait", "portrait", "Image", "image", "TokenImage", "tokenImage", "Token", "token")

        for slug, wrapper in self.entity_wrappers.items():
            try:
                items = wrapper.load_items()
            except Exception as exc:
                log_exception(
                    f"Unable to load items for '{slug}' while normalizing media paths: {exc}",
                    func_name="main_window.MainWindow._normalize_entity_media_paths",
                )
                continue

            updated = False
            for item in items:
                if not isinstance(item, dict):
                    continue

                for field in path_fields:
                    if field not in item:
                        continue

                    new_value, changed = self._normalize_single_media_path(item.get(field, ""), unique_candidates)
                    if changed:
                        item[field] = new_value
                        updated = True

            if updated:
                try:
                    wrapper.save_items(items)
                    log_info(
                        f"Normalized media paths for entity '{slug}'",
                        func_name="main_window.MainWindow._normalize_entity_media_paths",
                    )
                except Exception as exc:
                    log_exception(
                        f"Unable to save normalized media paths for '{slug}': {exc}",
                        func_name="main_window.MainWindow._normalize_entity_media_paths",
                    )

    def _normalize_single_media_path(self, value, base_candidates):
        """Return normalized (relative) path for portraits/images."""

        if not isinstance(value, str):
            return value, False

        raw = value
        stripped = raw.strip()
        if not stripped:
            return "", raw != ""

        normalized = stripped.replace("\\", "/")
        changed = normalized != raw

        if not self._is_absolute_path(normalized):
            if normalized.startswith("./"):
                normalized = normalized[2:]
                changed = True
            return normalized, changed

        path_obj = Path(normalized)

        for base in base_candidates:
            base_str = str(base).replace("\\", "/").rstrip("/")
            if not base_str:
                continue

            lower_value = normalized.lower()
            lower_base = base_str.lower()

            if lower_value == lower_base:
                return ".", True

            if lower_value.startswith(lower_base + "/"):
                relative_part = normalized[len(base_str) + 1 :]
                return relative_part, True

            try:
                rel_path = path_obj.relative_to(base)
                return rel_path.as_posix(), True
            except Exception:
                continue

        lowered = normalized.lower()
        for marker in ("/assets/", "/static/", "/images/"):
            idx = lowered.find(marker)
            if idx != -1:
                rel_path = normalized[idx + 1 :]
                return rel_path, True

        log_warning(
            f"Unable to derive relative path from '{value}', leaving as-is.",
            func_name="main_window.MainWindow._normalize_single_media_path",
        )
        return normalized, changed

    @staticmethod
    def _is_absolute_path(path_value: str) -> bool:
        if not path_value:
            return False

        if path_value.startswith(("/", "\\")):
            return True

        if re.match(r"^[a-zA-Z]:[\\/].*", path_value):
            return True

        if path_value.startswith("~"):
            return True

        return False

    def refresh_entities(self, *_):
        self.entity_definitions = load_entity_definitions()
        self.load_icons()
        self.init_wrappers()
        self.create_sidebar()


    def open_faction_graph_editor(self):
        self._graph_type = 'faction'
        self.current_gm_view = None
        self.clear_current_content()
        self.banner_toggle_btn.configure(state="normal")
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        editor = FactionGraphEditor(container, self.faction_wrapper)
        editor.pack(fill="both", expand=True)

        # keep both container and editor so we can snapshot/restore
        container.graph_editor = editor
        self.current_open_view   = container
        self.current_open_entity = None
        

    # =============================================================
    # Methods Called by Icon Buttons (Event Handlers)
    # =============================================================
    def clear_current_content(self):
        self._teardown_whiteboard_controller()
        # Always clear children of the inner content container
        try:
            for widget in self.inner_content_frame.winfo_children():
                widget.destroy()
        except Exception:
            pass

        if not self.banner_visible:
            # If banner is hidden, forget inner_content_frame from the grid so new content can occupy row 0
            try:
                self.inner_content_frame.grid_forget()
            except Exception:
                pass
            # Clear dynamic children of content_frame (exclude the static frames)
            for widget in self.content_frame.winfo_children():
                if widget not in (self.banner_frame, self.inner_content_frame):
                    try:
                        widget.destroy()
                    except Exception:
                        pass
        else:
            # If banner is visible, ensure inner_content_frame is at row 1 and ready for new content
            try:
                self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
            except Exception:
                pass

        # Leaving GM screen mode when clearing
        self._gm_mode = False
        self.current_gm_view = None

    def _teardown_whiteboard_controller(self):
        controller = getattr(self, "whiteboard_controller", None)
        if controller is None:
            return
        try:
            controller.close()
        except Exception:
            log_exception("Error while closing whiteboard", func_name="main_window.MainWindow._teardown_whiteboard_controller")
        finally:
            self.whiteboard_controller = None

    def move_current_view(self):
        """Move the current open view to the correct container based on banner state."""
        if self.current_open_view is not None:
            try:
                self.current_open_view.grid_forget()
            except tk.TclError:
                # If the widget was destroyed (because it was inside inner_content_frame), clear it
                self.current_open_view = None
                return

            parent = self.get_content_container()
            self.current_open_view.master = parent
            self.current_open_view.grid(row=0, column=0, sticky="nsew")

    def open_entity(self, entity):
        self.clear_current_content()
        target_parent = self.get_content_container()
        self.banner_toggle_btn._state="normal"
        container = ctk.CTkFrame(target_parent)
        container.grid(row=0, column=0, sticky="nsew")
        self.current_open_view = container
        self.current_open_entity = entity  # ✅ Add this clearly!

        wrapper = self.entity_wrappers.get(entity)
        if wrapper is None:
            wrapper = GenericModelWrapper(entity)
            self.entity_wrappers[entity] = wrapper

        template = load_template(entity)
        view = GenericListView(container, wrapper, template)
        view.pack(fill="both", expand=True)

        meta = self.entity_definitions.get(entity, {})
        display_label = meta.get("label") or entity.replace("_", " ").title()

        load_button = ctk.CTkButton(
            container,
            text=f"Load {display_label}",
            command=lambda: self.load_items_from_json(view, entity)
        )
        load_button.pack(side="right", padx=(5,5))
        # Assuming `editor_window` is your CTkToplevel or CTkFrame
        save_btn = ctk.CTkButton(
            container,
            text=f"Save {display_label}",
            command=lambda: self.save_items_to_json(view, entity)
        )
        save_btn.pack(side="right", padx=(5,5))

    def save_items_to_json(self, view, entity_name):
        display_label = self.entity_definitions.get(entity_name, {}).get(
            "label", entity_name.replace("_", " ").title()
        )
        # 1) Ask the user where to save
        path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")],
            title=f"Export {display_label} to JSON",
        )
        if not path:
            return  # user hit “Cancel”

        items = []
        try:
            with self._busy_operation(f"Saving {display_label}..."):
                try:
                    items = view.get_items()
                except AttributeError:
                    try:
                        items = view.items
                    except Exception:
                        wrapper = self.entity_wrappers.get(entity_name) or GenericModelWrapper(entity_name)
                        self.entity_wrappers[entity_name] = wrapper
                        items = wrapper.load_items()

                data = {entity_name: items}
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(data, f, indent=4)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save {display_label}:\n{e}")
            return

        messagebox.showinfo("Export Successful", f"Wrote {len(items)} {display_label} to:\n{path}")


    def load_items_from_json(self, view, entity_name):
        display_label = self.entity_definitions.get(entity_name, {}).get(
            "label", entity_name.replace("_", " ").title()
        )
        file_path = filedialog.askopenfilename(
            title=f"Load {display_label} from JSON",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        if not file_path:
            return
        items = []
        try:
            with self._busy_operation(f"Loading {display_label}..."):
                with open(file_path, "r", encoding="utf-8") as file:
                    data = json.load(file)
                items = data.get(entity_name, [])
                view.add_items(items)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load {display_label}: {e}")
            return

        messagebox.showinfo("Success", f"{len(items)} {display_label} loaded successfully!")

    def _auto_open_gm_screen_if_available(self):
        """Post-initialization hook to open the GM screen if scenarios exist."""
        scenario_wrapper = self.entity_wrappers.get("scenarios") or GenericModelWrapper("scenarios")
        self.entity_wrappers.setdefault("scenarios", scenario_wrapper)

        try:
            scenarios = scenario_wrapper.load_items()
        except Exception as exc:
            log_exception(
                f"Failed to load scenarios for initial GM Screen: {exc}",
                func_name="main_window.MainWindow._auto_open_gm_screen_if_available",
            )
            return

        if not scenarios:
            log_info(
                "Skipping automatic GM Screen launch because no scenarios exist",
                func_name="main_window.MainWindow._auto_open_gm_screen_if_available",
            )
            return

        self._prime_content_frames_for_gm_screen()
        self.open_gm_screen(show_empty_message=False)

    def open_gm_screen(self, *, show_empty_message=True):
        # 1) Clear any existing content
        self.clear_current_content()
        self._gm_mode = True
        # 2) Load all scenarios
        scenario_wrapper = self.entity_wrappers.get("scenarios") or GenericModelWrapper("scenarios")
        self.entity_wrappers.setdefault("scenarios", scenario_wrapper)
        scenarios = scenario_wrapper.load_items()
        if not scenarios:
            if show_empty_message:
                messagebox.showwarning("No Scenarios", "No scenarios available.")
            else:
                log_info(
                    "Skipped opening GM Screen because no scenarios are available",
                    func_name="main_window.MainWindow.open_gm_screen",
                )
            return

        layout_manager = GMScreenLayoutManager()
        layout_map = layout_manager.list_layouts()
        default_label = "Use Scenario Default"
        layout_options = [default_label]
        layout_options.extend(sorted(layout_map.keys()))
        selected_layout_var = tk.StringVar(value=default_label)

        # 3) Ensure the PC‐banner is shown and up to date
        if getattr(self, 'banner_frame', None) and self.banner_frame.winfo_exists():
            if not self.banner_frame.winfo_ismapped():
                self.banner_frame.pack(fill='x')
        else:
            # banner_frame was destroyed (or never created): re-create it
            self.banner_frame = self._create_banner_frame()
            self.banner_frame.pack(fill='x')
        pcs_items = {pc["Name"]: pc for pc in self.pc_wrapper.load_items()}
        if pcs_items:
            display_pcs_in_banner(self.banner_frame, pcs_items)

        # 4) Prepare inner content area
        self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
        for w in self.inner_content_frame.winfo_children():
            w.destroy()
        parent = self.inner_content_frame

        layout_bar = ctk.CTkFrame(parent)
        layout_bar.pack(fill="x", padx=10, pady=(5, 0))
        ctk.CTkLabel(layout_bar, text="Layout when opening scenario:").pack(side="left", padx=(0, 10), pady=5)
        layout_menu = ctk.CTkOptionMenu(layout_bar, variable=selected_layout_var, values=layout_options)
        layout_menu.pack(side="left", pady=5)

        # 5) Callback to open a selected scenario in detail
        def on_scenario_select(entity_type, entity_name):
            selected = next(
                (s for s in scenarios
                if s.get("Name", s.get("Title", "")) == entity_name),
                None
            )
            if not selected:
                messagebox.showwarning("Not Found", f"Scenario '{entity_name}' not found.")
                return
            # clear list and show scenario detail
            for w in parent.winfo_children():
                w.destroy()
            detail_container = ctk.CTkFrame(parent)
            detail_container.grid(row=0, column=0, sticky="nsew")
            chosen_layout = selected_layout_var.get()
            initial_layout = None if chosen_layout == default_label else chosen_layout
            view = GMScreenView(
                detail_container,
                scenario_item=selected,
                initial_layout=initial_layout,
                layout_manager=layout_manager,
            )
            view.pack(fill="both", expand=True)
            # track the active GM-screen view for our Ctrl+F handler
            self.current_gm_view = view

        # 6) Insert the generic list‐selection view
        list_selection = GenericListSelectionView(
            parent,
            "scenarios",
            scenario_wrapper,
            load_template("scenarios"),
            on_select_callback=on_scenario_select
        )
        list_selection.pack(fill="both", expand=True)
        self.current_gm_view = None
        # Configure banner and enable toggle in GM screen
        self.banner_visible = True
        try:
            self.banner_toggle_btn.configure(text="▲", state="normal")
            # Ensure underlying private state is also normalized where used
            self.banner_toggle_btn._state = "normal"
        except Exception:
            pass
        # 7) Lock banner and configure grid weights
        self.banner_visible = True
        self.banner_toggle_btn.configure(text="▲")
        self.banner_toggle_btn._state = "disabled"

        # Make row 0 (banner) fixed height, row 1 (content) expand
        self.content_frame.grid_rowconfigure(0, weight=0)
        self.content_frame.grid_rowconfigure(1, weight=1)
        self.content_frame.grid_columnconfigure(0, weight=1)

        # Make the inner_content_frame fully fill its cell
        self.inner_content_frame.grid_rowconfigure(0, weight=1)
        self.inner_content_frame.grid_columnconfigure(0, weight=1)
        # Ensure the banner toggle is enabled for GM screen after layout
        try:
            self.banner_toggle_btn.configure(text="▲", state="normal")
            self.banner_toggle_btn._state = "normal"
        except Exception:
            pass

    def open_npc_graph_editor(self):
        self._graph_type = 'npc'
        self.current_gm_view = None
        self.clear_current_content()
        self.banner_toggle_btn.configure(state="normal")
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        editor = NPCGraphEditor(container, self.npc_wrapper, self.faction_wrapper)
        editor.pack(fill="both", expand=True)

        # keep both container and editor so we can snapshot/restore
        container.graph_editor = editor
        self.current_open_view   = container
        self.current_open_entity = None


    def open_pc_graph_editor(self):
        self._graph_type = 'pc'
        self.current_gm_view = None
        self.clear_current_content()
        self.banner_toggle_btn.configure(state="normal")
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        editor = PCGraphEditor(container, self.pc_wrapper, self.faction_wrapper)
        editor.pack(fill="both", expand=True)

        container.graph_editor = editor
        self.current_open_view   = container
        self.current_open_entity = None


    def open_world_map(self):
        """Launch the World Map nested navigation window."""
        log_info("Opening World Map window", func_name="main_window.MainWindow.open_world_map")
        # If a GM screen is active, open as a tab instead
        try:
            if getattr(self, "current_gm_view", None) is not None and self.current_gm_view.winfo_exists():
                self.current_gm_view.open_world_map_tab()
                return
        except Exception:
            pass
        existing = getattr(self, "_world_map_window", None)
        if existing and existing.winfo_exists():
            existing.focus_force()
            existing.lift()
            existing.attributes("-topmost", True)
            existing.after_idle(lambda: existing.attributes("-topmost", False))
            return
        window = WorldMapWindow(self)
        window.lift()
        window.focus_force()
        window.attributes("-topmost", True)
        window.after_idle(lambda: window.attributes("-topmost", False))
        self._world_map_window = window

    def open_scenario_graph_editor(self):
        self._graph_type = 'scenario'
        self.current_gm_view = None
        self.clear_current_content()
        self.banner_toggle_btn.configure(state="normal")
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        editor = ScenarioGraphEditor(
            container,
            GenericModelWrapper("scenarios"),
            GenericModelWrapper("npcs"),
            GenericModelWrapper("creatures"),
            GenericModelWrapper("places")
        )
        editor.pack(fill="both", expand=True)

        container.graph_editor = editor
        self.current_open_view   = container
        self.current_open_entity = None

    def open_scene_flow_viewer(self):
        from modules.scenarios.scene_flow_viewer import SceneFlowViewerWindow
        # If a GM screen is active, open as a tab instead
        try:
            if getattr(self, "current_gm_view", None) is not None and self.current_gm_view.winfo_exists():
                self.current_gm_view.open_scene_flow_tab()
                return
        except Exception:
            pass
        if getattr(self, "_scene_flow_window", None) and self._scene_flow_window.winfo_exists():
            try:
                self._scene_flow_window.focus()
                self._scene_flow_window.lift()
            except Exception:
                pass
            return

        def _on_close():
            self._scene_flow_window = None

        window = SceneFlowViewerWindow(
            self,
            scenario_wrapper=GenericModelWrapper("scenarios"),
            npc_wrapper=GenericModelWrapper("npcs"),
            creature_wrapper=GenericModelWrapper("creatures"),
            place_wrapper=GenericModelWrapper("places"),
            on_close=_on_close,
        )
        self._scene_flow_window = window

    def export_foundry(self):
        preview_and_export_foundry(self)

    def _queue_update_check(self, *, force: bool = False):
        try:
            updates_enabled = ConfigHelper.getboolean("Updates", "enabled", fallback=True)
        except Exception:
            updates_enabled = True
        if not updates_enabled:
            return
        if self._update_thread and self._update_thread.is_alive():
            return
        self._update_thread = threading.Thread(
            target=self._async_check_for_updates,
            kwargs={"force": force},
            daemon=True,
        )
        self._update_thread.start()

    def _async_check_for_updates(self, *, force: bool = False):
        channel = ConfigHelper.get("Updates", "channel", fallback="stable") or "stable"
        preferred_asset = ConfigHelper.get("Updates", "asset_name", fallback="") or None
        try:
            interval_hours = float(
                ConfigHelper.get("Updates", "check_interval_hours", fallback="24") or "24"
            )
        except (TypeError, ValueError):
            interval_hours = 24.0
        last_check_raw = ConfigHelper.get("Updates", "last_check", fallback="") or ""
        if last_check_raw and not force:
            try:
                last_check_ts = float(last_check_raw)
            except (TypeError, ValueError):
                last_check_ts = 0.0
            else:
                if time.time() - last_check_ts < interval_hours * 3600:
                    log_debug(
                        "Skipping update check; interval not yet elapsed",
                        func_name="main_window.MainWindow._async_check_for_updates",
                    )
                    return
        try:
            current_version, candidate = update_helper.check_for_update(
                channel=channel,
                preferred_asset=preferred_asset,
            )
            ConfigHelper.set("Updates", "last_check", str(time.time()))
            ConfigHelper.load_config()
        except Exception as exc:
            log_warning(
                f"Update check failed: {exc}",
                func_name="main_window.MainWindow._async_check_for_updates",
            )
            return

        if not candidate:
            return

        log_info(
            f"Update available: {candidate.version}",
            func_name="main_window.MainWindow._async_check_for_updates",
        )
        self.after(0, lambda: self._prompt_update(str(current_version), candidate))

    def _prompt_update(self, current_version: str, candidate: update_helper.UpdateCandidate):
        release_notes = (candidate.release_notes or "").strip()
        if len(release_notes) > 800:
            release_notes = release_notes[:800].rstrip() + "\u2026"
        message_lines = [
            f"A new version ({candidate.version}) is available.",
            f"You are currently running {current_version}.",
        ]
        if release_notes:
            message_lines.append("Release notes:")
            message_lines.append(release_notes)
        message_lines.append("Download and install now?")

        if not messagebox.askyesno("Update Available", "\n\n".join(message_lines)):
            return

        self._begin_update_download(candidate)

    def _begin_update_download(self, candidate: update_helper.UpdateCandidate):
        log_info(
            f"Preparing download for update {candidate.version}",
            func_name="main_window.MainWindow._begin_update_download",
        )
        if getattr(sys, "frozen", False):
            install_root = Path(sys.executable).resolve().parent
            restart_target = sys.executable
        else:
            install_root = Path(__file__).resolve().parent
            restart_target = None
        preserve = ["Campaigns", "config/config.ini"]

        def worker(progress_cb):
            stage_root, payload_root = update_helper.prepare_staging_area(
                candidate,
                progress_callback=progress_cb,
            )
            process = update_helper.launch_installer(
                payload_root,
                install_root=install_root,
                restart_target=restart_target,
                wait_for_pid=os.getpid(),
                preserve=preserve,
                cleanup_root=stage_root,
            )
            return {"installer_pid": getattr(process, "pid", None)}

        def detail_builder(result):
            pid = result.get("installer_pid") if isinstance(result, dict) else None
            detail = (
                "An installer helper is waiting for GMCampaignDesigner to close before copying the new files."
            )
            if pid:
                detail += f" (PID {pid})"
            detail += " Close the app when you are ready so the update can finish."
            return detail

        self._run_progress_task(
            "Downloading Update",
            worker,
            "Update downloaded.",
            detail_builder=detail_builder,
        )

    def _run_progress_task(self, title, worker, success_message, detail_builder=None):
        progress_win = ctk.CTkToplevel(self)
        progress_win.title(title)
        progress_win.geometry("420x180")
        progress_win.resizable(False, False)
        progress_win.transient(self)
        progress_win.grab_set()
        progress_win.lift()
        progress_win.focus_force()
        progress_win.protocol("WM_DELETE_WINDOW", lambda: None)

        progress_label = ctk.CTkLabel(progress_win, text="Starting...", wraplength=360, justify="center")
        progress_label.pack(fill="x", padx=20, pady=(20, 10))
        progress_bar = ctk.CTkProgressBar(progress_win, mode="determinate")
        progress_bar.pack(fill="x", padx=20, pady=(0, 20))
        progress_bar.set(0.0)

        def update(message: str, fraction: float) -> None:
            def _apply():
                progress_label.configure(text=message)
                try:
                    progress_bar.set(max(0.0, min(1.0, float(fraction))))
                except Exception:
                    progress_bar.set(0.0)

            self.after(0, _apply)

        def close_window():
            if progress_win.winfo_exists():
                try:
                    progress_win.grab_release()
                except Exception:
                    pass
                progress_win.destroy()

        def handle_error(title: str, detail: str) -> None:
            close_window()
            messagebox.showerror(title, detail)

        def on_success(result):
            close_window()
            detail = detail_builder(result) if detail_builder else None
            message = success_message or "Operation completed."
            if detail:
                message = f"{message}\n\n{detail}"
            messagebox.showinfo("Success", message)

        def run_worker():
            try:
                result = worker(update)
            except PermissionError as exc:
                detail = str(exc)
                self.after(0, lambda detail=detail: handle_error("Permission Denied", detail))
                return
            except ManifestError as exc:
                detail = str(exc)
                self.after(0, lambda detail=detail: handle_error("Invalid Backup", detail))
                return
            except BackupError as exc:
                detail = str(exc)
                self.after(0, lambda detail=detail: handle_error("Backup Error", detail))
                return
            except Exception as exc:
                log_exception(
                    f"Unexpected error during {title}: {exc}",
                    func_name="main_window.MainWindow._run_progress_task",
                )
                detail = str(exc)
                self.after(0, lambda detail=detail: handle_error("Unexpected Error", detail))
                return

            self.after(0, lambda: on_success(result))

        threading.Thread(target=run_worker, daemon=True).start()

    def _format_backup_summary(self, manifest: dict | None, *, include_target: bool) -> str:
        if not manifest:
            return ""

        lines: list[str] = []
        campaign_name = manifest.get("campaign_name")
        if campaign_name:
            lines.append(f"Campaign: {campaign_name}")

        created = manifest.get("created_at")
        if created:
            lines.append(f"Created: {created}")

        archive_path = manifest.get("archive_path")
        if archive_path:
            lines.append(f"Archive: {archive_path}")

        if include_target:
            target = manifest.get("restored_to")
            if target:
                lines.append(f"Restored to: {target}")

        files = manifest.get("files")
        if isinstance(files, (list, tuple)):
            lines.append(f"Files included: {len(files)}")

        missing = manifest.get("missing") or []
        if missing:
            lines.append(f"Missing at creation: {len(missing)}")

        return "\n".join(lines)

    def prompt_campaign_backup(self):
        campaign_dir = Path(ConfigHelper.get_campaign_dir()).resolve()
        if not campaign_dir.exists():
            messagebox.showerror(
                "Campaign Directory Missing",
                f"The configured campaign directory was not found:\n{campaign_dir}",
            )
            return

        timestamp = time.strftime("%Y%m%d_%H%M%S")
        default_name = f"{campaign_dir.name or 'campaign'}_{timestamp}.zip"
        destination = filedialog.asksaveasfilename(
            title="Create Campaign Backup",
            initialdir=str(campaign_dir),
            initialfile=default_name,
            defaultextension=".zip",
            filetypes=[("Zip Archives", "*.zip")],
        )
        if not destination:
            return

        if os.path.exists(destination) and not messagebox.askyesno(
            "Overwrite Existing File",
            f"The selected file already exists:\n{destination}\n\nOverwrite it?",
        ):
            return

        if not messagebox.askyesno(
            "Confirm Backup",
            f"Create a backup archive at:\n{destination}?",
        ):
            return

        self._run_progress_task(
            "Creating Backup",
            lambda cb: create_backup_archive(destination, cb),
            "Backup archive created successfully.",
            lambda manifest: self._format_backup_summary(manifest, include_target=False),
        )

    @staticmethod
    def _sanitize_campaign_name(name: str) -> str:
        safe = "".join(ch for ch in name.strip() if ch.isalnum() or ch in ("_", "-", " "))
        safe = safe.strip().replace(" ", "_")
        return safe

    def prompt_campaign_restore(self):
        campaign_dir = Path(ConfigHelper.get_campaign_dir()).resolve()
        if campaign_dir.exists():
            initial_dir = campaign_dir
        else:
            try:
                initial_dir = Path.home()
            except Exception:
                initial_dir = Path.cwd()
        archive = filedialog.askopenfilename(
            title="Select Backup Archive",
            initialdir=str(initial_dir),
            filetypes=[("Zip Archives", "*.zip"), ("All Files", "*.*")],
        )
        if not archive:
            return

        try:
            manifest = read_backup_manifest(archive)
        except ManifestError as exc:
            messagebox.showerror("Invalid Backup", str(exc))
            return
        except BackupError as exc:
            messagebox.showerror("Backup Error", str(exc))
            return

        default_name = manifest.get("campaign_name") or ""
        if not default_name:
            db_path = manifest.get("database_path")
            if isinstance(db_path, str) and db_path:
                default_name = Path(db_path).stem
            else:
                default_name = Path(archive).stem

        base_dir = campaign_dir.parent if campaign_dir.parent != campaign_dir else campaign_dir

        while True:
            new_name = simpledialog.askstring(
                "Restore Campaign",
                "Enter a name for the restored campaign:",
                initialvalue=default_name,
                parent=self,
            )
            if new_name is None:
                return
            cleaned_name = self._sanitize_campaign_name(new_name)
            if not cleaned_name:
                messagebox.showwarning("Invalid Name", "Please enter a valid campaign name.")
                continue
            destination_dir = base_dir / cleaned_name
            db_filename = f"{cleaned_name}.db"
            display_name = new_name.strip() or cleaned_name
            break

        if destination_dir.exists():
            if not messagebox.askyesno(
                "Confirm Overwrite",
                "The selected campaign name already exists:\n"
                f"{destination_dir}\n\nRestoring will overwrite files in this directory. Continue?",
            ):
                return

        if not messagebox.askyesno(
            "Confirm Restore",
            "Restoring a backup will overwrite files in:\n"
            f"{destination_dir}\n\nArchive:\n{archive}\n\nProceed?",
        ):
            return

        destination_dir.mkdir(parents=True, exist_ok=True)
        new_db_path = destination_dir / db_filename

        def detail_builder(manifest_data: dict | None) -> str:
            if manifest_data:
                ConfigHelper.set("Database", "path", str(new_db_path))
            return self._format_backup_summary(manifest_data, include_target=True)

        self._run_progress_task(
            "Restoring Backup",
            lambda cb: restore_backup_archive(
                archive,
                destination_dir,
                cb,
                campaign_name=display_name,
                database_filename=db_filename,
            ),
            "Backup restored successfully.",
            detail_builder,
        )

    def open_scenario_importer(self):
        self.clear_current_content()
        container = ctk.CTkFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew")
        ScenarioImportWindow(container)

    def _on_scenario_built(self):
        try:
            if self.current_open_entity == "scenarios" and self.current_open_view is not None:
                for child in self.current_open_view.winfo_children():
                    if isinstance(child, GenericListView):
                        child.reload_from_db()
                        break
        except Exception as exc:
            log_exception(
                f"Failed to refresh scenarios after builder save: {exc}",
                func_name="main_window.MainWindow._on_scenario_built",
            )

    def open_creature_importer(self):
        from modules.creatures.creature_importer import CreatureImportWindow

        self.clear_current_content()
        container = ctk.CTkFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew")
        CreatureImportWindow(container)

    def open_object_importer(self):
        self.clear_current_content()
        container = ctk.CTkFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew")
        ObjectImportWindow(container)

    def open_scenario_generator(self):
        self.clear_current_content()
        parent = self.get_content_container()
        container = ScenarioGeneratorView(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)
        self.current_open_view = container
        self.current_open_entity = None

    def open_scenario_builder(self):
        try:
            wizard = ScenarioBuilderWizard(self, on_saved=self._on_scenario_built)
            wizard.grab_set()
            wizard.focus_force()
        except Exception as exc:
            log_exception(
                f"Failed to open Scenario Builder Wizard: {exc}",
                func_name="main_window.MainWindow.open_scenario_builder",
            )
            messagebox.showerror("Error", f"Failed to open Scenario Builder:\n{exc}")

    def open_random_table_editor(self):
        try:
            dialog = RandomTableEditorDialog(self)
            dialog.grab_set()
            dialog.focus_force()
        except Exception as exc:
            log_exception(
                f"Failed to open Random Table Editor: {exc}",
                func_name="main_window.MainWindow.open_random_table_editor",
            )
            messagebox.showerror("Random Tables", f"Unable to open editor:\n{exc}")

    def change_database_storage(self):
        current_path = ConfigHelper.get("Database", "path", fallback="") or None

        if self._database_manager_dialog is not None:
            try:
                self._database_manager_dialog.lift()
                self._database_manager_dialog.focus_force()
                return
            except Exception:
                self._database_manager_dialog = None

        def _on_selected(path: str, is_new: bool) -> None:
            self._database_manager_dialog = None
            self._apply_database_selection(path, is_new)

        def _on_cancelled() -> None:
            self._database_manager_dialog = None

        dialog = DatabaseManagerDialog(
            self,
            current_path=current_path,
            on_selected=_on_selected,
            on_cancelled=_on_cancelled,
        )
        self._database_manager_dialog = dialog

    def _apply_database_selection(self, new_db_path: str, is_new_db: bool) -> None:
        if not new_db_path:
            return

        normalized_path = os.path.abspath(os.path.normpath(new_db_path))
        try:
            os.makedirs(os.path.dirname(normalized_path), exist_ok=True)
        except Exception:
            log_exception(
                f"Failed to prepare campaign directory for {normalized_path}",
                func_name="MainWindow._apply_database_selection",
            )

        ConfigHelper.set("Database", "path", normalized_path)

        if is_new_db:
            try:
                from shutil import copyfile

                entities = (
                    "pcs",
                    "npcs",
                    "scenarios",
                    "factions",
                    "places",
                    "objects",
                    "creatures",
                    "informations",
                    "clues",
                    "maps",
                    "books",
                )
                camp_dir = os.path.abspath(os.path.dirname(normalized_path))
                tpl_dir = os.path.join(camp_dir, "templates")
                os.makedirs(tpl_dir, exist_ok=True)
                for entity in entities:
                    src = os.path.join("modules", entity, f"{entity}_template.json")
                    dst = os.path.join(tpl_dir, f"{entity}_template.json")
                    try:
                        if not os.path.exists(dst):
                            copyfile(src, dst)
                    except Exception:
                        pass
            except Exception:
                log_exception(
                    "Failed to seed default templates for new campaign.",
                    func_name="MainWindow._apply_database_selection",
                )

        initialize_db()
        self._reload_active_campaign_system()

        conn = sqlite3.connect(normalized_path)
        cursor = conn.cursor()

        for entity in load_entity_definitions().keys():
            ensure_entity_schema(entity)

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS nodes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                npc_name TEXT,
                x INTEGER,
                y INTEGER,
                color TEXT
            )
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS links (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                npc_name1 TEXT,
                npc_name2 TEXT,
                text TEXT,
                arrow_mode TEXT
            )
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS shapes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                type TEXT,
                x INTEGER,
                y INTEGER,
                w INTEGER,
                h INTEGER,
                color TEXT,
                tag TEXT,
                z INTEGER
            )
        """)

        conn.commit()
        conn.close()

        self.refresh_entities()

        db_name = os.path.splitext(os.path.basename(normalized_path))[0]
        self.db_name_label.configure(text=db_name)
        try:
            full_path = os.path.abspath(normalized_path)
            if getattr(self, "db_tooltip", None) is None:
                self.db_tooltip = ToolTip(self.db_name_label, full_path)
            else:
                self.db_tooltip.text = full_path
        except Exception:
            self.db_tooltip = None

    def select_swarmui_path(self):
        folder = filedialog.askdirectory(title="Select SwarmUI Path")
        if folder:
            ConfigHelper.set("Paths", "swarmui_path", folder)
            messagebox.showinfo("SwarmUI Path Set", f"SwarmUI path set to:\n{folder}")

    def launch_swarmui(self):
        global SWARMUI_PROCESS
        swarmui_path = ConfigHelper.get("Paths", "swarmui_path", fallback=r"E:\SwarmUI\SwarmUI")
        SWARMUI_CMD = os.path.join(swarmui_path, "launch-windows.bat")
        env = os.environ.copy()
        env.pop('VIRTUAL_ENV', None)
        if SWARMUI_PROCESS is None or SWARMUI_PROCESS.poll() is not None:
            try:
                SWARMUI_PROCESS = subprocess.Popen(
                    SWARMUI_CMD,
                    shell=True,
                    cwd=swarmui_path,
                    env=env
                )
                time.sleep(20.0)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to launch SwarmUI: {e}")

    def cleanup_swarmui(self):
        global SWARMUI_PROCESS
        if SWARMUI_PROCESS is not None and SWARMUI_PROCESS.poll() is None:
            SWARMUI_PROCESS.terminate()

    # ------------------------------------------------------
    # Unified Generate Portraits for NPCs and Creatures
    # ------------------------------------------------------
    def generate_missing_portraits(self):
        top = ctk.CTkToplevel(self)
        top.title("Generate Portraits")
        top.geometry("300x150")
        top.transient(self)
        top.grab_set()
        # Use ctk.StringVar (not CTkStringVar)
        selection = ctk.StringVar(value="NPC")
        ctk.CTkLabel(top, text="Generate portraits for:").pack(pady=10)
        ctk.CTkRadioButton(top, text="NPCs", variable=selection, value="NPC").pack(pady=5)
        ctk.CTkRadioButton(top, text="Creatures", variable=selection, value="Creature").pack(pady=5)
        def on_confirm():
            choice = selection.get()
            top.destroy()
            if choice == "NPC":
                self.generate_missing_npc_portraits()
            else:
                self.generate_missing_creature_portraits()
        ctk.CTkButton(top, text="Continue", command=on_confirm).pack(pady=10)

    def generate_missing_npc_portraits(self):
        def confirm_model_and_continue():
            ConfigHelper.set("LastUsed", "model", self.selected_model.get())
            top.destroy()
            self.generate_portraits_continue_npcs()
        top = ctk.CTkToplevel(self)
        top.title("Select AI Model for NPCs")
        top.geometry("400x200")
        top.transient(self)
        top.grab_set()
        ctk.CTkLabel(top, text="Select AI Model to use for NPC portrait generation:").pack(pady=20)
        last_model = ConfigHelper.get("LastUsed", "model", fallback=None)
        # Use ctk.StringVar
        if last_model in self.model_options:
            self.selected_model = ctk.StringVar(value=last_model)
        else:
            self.selected_model = ctk.StringVar(value=self.model_options[0])
        ctk.CTkOptionMenu(top, values=self.model_options, variable=self.selected_model).pack(pady=10)
        ctk.CTkButton(top, text="Continue", command=confirm_model_and_continue).pack(pady=10)

    def generate_portraits_continue_npcs(self):
        db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM npcs")
        npc_rows = cursor.fetchall()
        modified = False
        for npc in npc_rows:
            portrait = npc["Portrait"] if npc["Portrait"] is not None else ""
            if not portrait.strip():
                npc_dict = dict(npc)
                self.generate_portrait_for_npc(npc_dict)
                if npc_dict.get("Portrait"):
                    cursor.execute("UPDATE npcs SET Portrait = ? WHERE Name = ?", (npc_dict["Portrait"], npc["Name"]))
                    modified = True
        if modified:
            conn.commit()
            print("Updated NPC database with generated portraits.")
        else:
            print("No NPCs were missing portraits.")
        conn.close()

    def generate_missing_creature_portraits(self):
        def confirm_model_and_continue():
            ConfigHelper.set("LastUsed", "model", self.selected_model.get())
            top.destroy()
            self.generate_portraits_continue_creatures()
        top = ctk.CTkToplevel(self)
        top.title("Select AI Model for Creatures")
        top.geometry("400x200")
        top.transient(self)
        top.grab_set()
        ctk.CTkLabel(top, text="Select AI Model to use for creature portrait generation:").pack(pady=20)
        last_model = ConfigHelper.get("LastUsed", "model", fallback=None)
        if last_model in self.model_options:
            self.selected_model = ctk.StringVar(value=last_model)
        else:
            self.selected_model = ctk.StringVar(value=self.model_options[0])
        ctk.CTkOptionMenu(top, values=self.model_options, variable=self.selected_model).pack(pady=10)
        ctk.CTkButton(top, text="Continue", command=confirm_model_and_continue).pack(pady=10)

    def generate_portraits_continue_creatures(self):
        db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM creatures")
        creature_rows = cursor.fetchall()
        modified = False
        for creature in creature_rows:
            portrait = creature["Portrait"] if creature["Portrait"] is not None else ""
            if not portrait.strip():
                creature_dict = dict(creature)
                self.generate_portrait_for_creature(creature_dict)
                if creature_dict.get("Portrait"):
                    cursor.execute("UPDATE creatures SET Portrait = ? WHERE Name = ?", (creature_dict["Portrait"], creature_dict["Name"]))
                    modified = True
        if modified:
            conn.commit()
            print("Updated creature database with generated portraits.")
        else:
            print("No creatures were missing portraits.")
        conn.close()

    def generate_portrait_for_npc(self, npc):
        self.launch_swarmui()
        SWARM_API_URL = "http://127.0.0.1:7801"
        try:
            session_url = f"{SWARM_API_URL}/API/GetNewSession"
            session_response = requests.post(session_url, json={}, headers={"Content-Type": "application/json"})
            session_data = session_response.json()
            session_id = session_data.get("session_id")
            if not session_id:
                print(f"Failed to obtain session ID for NPC {npc.get('Name', 'Unknown')}")
                return
            npc_name = npc.get("Name", "Unknown")
            npc_role = npc.get("Role", "Unknown")
            npc_faction = npc.get("Factions", "Unknown")
            npc_desc = npc.get("Description", "Unknown")
            npc_desc = text_helpers.format_longtext(npc_desc)
            prompt = f"{npc_name} {npc_desc} {npc_role} {npc_faction}"
            prompt_data = {
                "session_id": session_id,
                "images": 1,
                "prompt": prompt,
                "negativeprompt": ("blurry, low quality, comics style, mangastyle, paint style, watermark, ugly, "
                                "monstrous, too many fingers, too many legs, too many arms, bad hands, "
                                "unrealistic weapons, bad grip on equipment, nude"),
                "model": self.selected_model.get(),
                "width": 1024,
                "height": 1024,
                "cfgscale": 9,
                "steps": 20,
                "seed": -1
            }
            generate_url = f"{SWARM_API_URL}/API/GenerateText2Image"
            image_response = requests.post(generate_url, json=prompt_data, headers={"Content-Type": "application/json"})
            image_data = image_response.json()
            images = image_data.get("images")
            if not images or len(images) == 0:
                print(f"Image generation failed for NPC '{npc_name}'")
                return
            image_url = f"{SWARM_API_URL}/{images[0]}"
            downloaded_image = requests.get(image_url)
            if downloaded_image.status_code != 200:
                print(f"Failed to download generated image for NPC '{npc_name}'")
                return
            output_filename = f"{npc_name.replace(' ', '_')}_portrait.png"
            with open(output_filename, "wb") as f:
                f.write(downloaded_image.content)
            GENERATED_FOLDER = os.path.join(ConfigHelper.get_campaign_dir(), "assets", "generated")
            os.makedirs(GENERATED_FOLDER, exist_ok=True)
            shutil.copy(output_filename, os.path.join(GENERATED_FOLDER, output_filename))
            npc["Portrait"] = self.copy_and_resize_portrait(npc, output_filename)
            os.remove(output_filename)
            print(f"Generated portrait for NPC '{npc_name}'")
        except Exception as e:
            print(f"Error generating portrait for NPC '{npc.get('Name', 'Unknown')}': {e}")

    def generate_portrait_for_creature(self, creature):
        self.launch_swarmui()
        SWARM_API_URL = "http://127.0.0.1:7801"
        try:
            session_url = f"{SWARM_API_URL}/API/GetNewSession"
            session_response = requests.post(session_url, json={}, headers={"Content-Type": "application/json"})
            session_data = session_response.json()
            session_id = session_data.get("session_id")
            if not session_id:
                print(f"Failed to obtain session ID for Creature {creature.get('Name', 'Unknown')}")
                return
            creature_name = creature.get("Name", "Unknown")
            creature_desc = creature.get("Description", "Unknown")
            stats = creature.get("Stats", "")
            creature_desc_formatted = text_helpers.format_longtext(creature_desc)
            prompt = f"{creature_name} {creature_desc_formatted} {stats}"
            prompt_data = {
                "session_id": session_id,
                "images": 1,
                "prompt": prompt,
                "negativeprompt": ("blurry, low quality, comics style, mangastyle, paint style, watermark, ugly, "
                                "monstrous, too many fingers, too many legs, too many arms, bad hands, "
                                "unrealistic weapons, bad grip on equipment, nude"),
                "model": self.selected_model.get(),
                "width": 1024,
                "height": 1024,
                "cfgscale": 9,
                "steps": 20,
                "seed": -1
            }
            generate_url = f"{SWARM_API_URL}/API/GenerateText2Image"
            image_response = requests.post(generate_url, json=prompt_data, headers={"Content-Type": "application/json"})
            image_data = image_response.json()
            images = image_data.get("images")
            if not images or len(images) == 0:
                print(f"Image generation failed for Creature '{creature_name}'")
                return
            image_url = f"{SWARM_API_URL}/{images[0]}"
            downloaded_image = requests.get(image_url)
            if downloaded_image.status_code != 200:
                print(f"Failed to download generated image for Creature '{creature_name}'")
                return
            output_filename = f"{creature_name.replace(' ', '_')}_portrait.png"
            with open(output_filename, "wb") as f:
                f.write(downloaded_image.content)
            GENERATED_FOLDER = os.path.join(ConfigHelper.get_campaign_dir(), "assets", "generated")
            os.makedirs(GENERATED_FOLDER, exist_ok=True)
            shutil.copy(output_filename, os.path.join(GENERATED_FOLDER, output_filename))
            creature["Portrait"] = self.copy_and_resize_portrait(creature, output_filename)
            os.remove(output_filename)
            print(f"Generated portrait for Creature '{creature_name}'")
        except Exception as e:
            print(f"Error generating portrait for Creature '{creature.get('Name', 'Unknown')}': {e}")

    def copy_and_resize_portrait(self, entity, src_path):
        campaign_dir = ConfigHelper.get_campaign_dir()
        PORTRAIT_FOLDER = os.path.join(campaign_dir, "assets", "portraits")
        MAX_PORTRAIT_SIZE = (1024, 1024)
        os.makedirs(PORTRAIT_FOLDER, exist_ok=True)
        name = entity.get("Name", "Unnamed").replace(" ", "_")
        ext = os.path.splitext(src_path)[-1].lower()
        dest_filename = f"{name}_{id(self)}{ext}"
        dest_path = os.path.join(PORTRAIT_FOLDER, dest_filename)
        with Image.open(src_path) as img:
            img = img.convert("RGB")
            img.thumbnail(MAX_PORTRAIT_SIZE)
            img.save(dest_path)
        try:
            relative_path = os.path.relpath(dest_path, campaign_dir)
        except ValueError:
            return dest_path
        if relative_path.startswith(".."):
            return dest_path
        return relative_path.replace(os.sep, "/")

    def import_portraits_from_directory(self):
        """Match and import portraits from a directory for all portrait-capable entities."""

        self.portrait_importer.import_portraits_from_directory()

    def preview_and_export_scenarios(self):
        scenario_wrapper = GenericModelWrapper("scenarios")
        scenario_items = scenario_wrapper.load_items()
        if not scenario_items:
            messagebox.showwarning("No Scenarios", "There are no scenarios available.")
            return
        selection_window = Toplevel()
        selection_window.title("Select Scenarios to Export")
        selection_window.geometry("400x300")
        listbox = Listbox(selection_window, selectmode="multiple", height=15)
        listbox.pack(fill="both", expand=True, padx=10, pady=10)
        scenario_titles = [scenario.get("Title", "Unnamed Scenario") for scenario in scenario_items]
        for title in scenario_titles:
            listbox.insert("end", title)
        def export_selected():
            selected_indices = listbox.curselection()
            if not selected_indices:
                messagebox.showwarning("No Selection", "Please select at least one scenario to export.")
                return
            selected_scenarios = [scenario_items[i] for i in selected_indices]
            self.preview_and_save(selected_scenarios)
            selection_window.destroy()
        ctk.CTkButton(selection_window, text="Export Selected", command=export_selected).pack(pady=5)

    def preview_and_save(self, selected_scenarios):
        creature_items = {creature["Name"]: creature for creature in self.creature_wrapper.load_items()}
        place_items = {place["Name"]: place for place in self.place_wrapper.load_items()}
        npc_items = {npc["Name"]: npc for npc in self.npc_wrapper.load_items()}

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Files", "*.docx"), ("All Files", "*.*")],
            title="Save Scenario Export"
        )
        if not file_path:
            return

        doc = Document()
        doc.add_heading("Campaign Scenarios", level=1)
        for scenario in selected_scenarios:
            title = scenario.get("Title", "Unnamed Scenario")
            summary = scenario.get("Summary", "No description provided.")
            secrets = scenario.get("Secrets", "No secrets provided.")

            doc.add_heading(title, level=2)
            doc.add_heading("Summary", level=3)
            if isinstance(summary, dict):
                p = doc.add_paragraph()
                run = p.add_run(summary.get("text", ""))
                self.apply_formatting(run, summary.get("formatting", {}))
            else:
                doc.add_paragraph(str(summary))

            doc.add_heading("Secrets", level=3)
            if isinstance(secrets, dict):
                p = doc.add_paragraph()
                run = p.add_run(secrets.get("text", ""))
                self.apply_formatting(run, secrets.get("formatting", {}))
            else:
                doc.add_paragraph(str(secrets))
            scenes = scenario.get("Scenes") or []
            if scenes:
                doc.add_heading("Scenes", level=3)
                for scene in scenes:
                    scene_title = ""
                    text_payload = scene
                    links_payload = []
                    if isinstance(scene, dict):
                        scene_title = scene.get("Title") or scene.get("Scene") or ""
                        text_payload = scene.get("Text") or scene.get("text") or scene
                        links_payload = scene.get("Links") or []
                    if scene_title:
                        doc.add_paragraph(scene_title, style="List Bullet")
                    if isinstance(text_payload, dict):
                        p = doc.add_paragraph()
                        run = p.add_run(text_payload.get("text", ""))
                        self.apply_formatting(run, text_payload.get("formatting", {}))
                    else:
                        doc.add_paragraph(str(text_payload))
                    if links_payload:
                        for link in links_payload:
                            target = link.get("Target") or link.get("target") or ""
                            text = link.get("Text") or link.get("text") or "Continue"
                            doc.add_paragraph(f"  → {text}: {target}")
            
            # Places Section
            doc.add_heading("Places", level=3)
            for place_name in scenario.get("Places", []):
                place = place_items.get(place_name, {"Name": place_name, "Description": "Unknown Place"})
                if isinstance(place["Description"], dict):
                    description_text = place["Description"].get("text", "Unknown Place")
                else:
                    description_text = place["Description"]
                doc.add_paragraph(f"- {place['Name']}: {description_text}")

            # NPCs Section
            doc.add_heading("NPCs", level=3)
            for npc_name in scenario.get("NPCs", []):
                npc = npc_items.get(npc_name, {"Name": npc_name, "Role": "Unknown",
                                                "Description": {"text": "Unknown NPC", "formatting": {}}})
                p = doc.add_paragraph(f"- {npc['Name']} ({npc['Role']}, {npc.get('Faction', 'Unknown')}): ")
                description = npc['Description']
                if isinstance(description, dict):
                    run = p.add_run(description.get("text", ""))
                    self.apply_formatting(run, description.get("formatting", {}))
                else:
                    p.add_run(str(description))

            # Creatures Section
            doc.add_heading("Creatures", level=3)
            for creature_name in scenario.get("Creatures", []):
                creature = creature_items.get(creature_name, {
                    "Name": creature_name,
                    "Stats": {"text": "No Stats", "formatting": {}},
                    "Powers": {"text": "No Powers", "formatting": {}},
                    "Description": {"text": "Unknown Creature", "formatting": {}}
                })
                stats = creature["Stats"]
                if isinstance(stats, dict):
                    stats_text = stats.get("text", "No Stats")
                else:
                    stats_text = stats
                powers = creature.get("Powers", "Unknown")
                if isinstance(powers, dict):
                    powers_text = powers.get("text", "No Powers")
                else:
                    powers_text = powers
                p = doc.add_paragraph(f"- {creature['Name']} ({stats_text}, {powers_text}): ")
                description = creature["Description"]
                if isinstance(description, dict):
                    run = p.add_run(description.get("text", ""))
                    self.apply_formatting(run, description.get("formatting", {}))
                else:
                    p.add_run(str(description))
        doc.save(file_path)
        messagebox.showinfo("Export Successful", f"Scenario exported successfully to:\n{file_path}")

    def apply_formatting(self, run, formatting):
        if formatting.get('bold'):
            run.bold = True
        if formatting.get('italic'):
            run.italic = True
        if formatting.get('underline'):
            run.underline = True

    def normalize_name(self, name):
        if name is None:
            return ""
        text = unicodedata.normalize("NFKD", str(name))
        text = "".join(ch for ch in text if not unicodedata.combining(ch))
        text = text.lower().replace('_', ' ')
        text = re.sub(r"[^a-z0-9]+", " ", text)
        return " ".join(text.split())

    def build_portrait_mapping(self):
        mapping = {}
        campaign_dir = ConfigHelper.get_campaign_dir()
        dir_txt_path = os.path.join(campaign_dir, "assets", "portraits", "dir.txt")
        if not os.path.exists(dir_txt_path):
            print(f"dir.txt not found at {dir_txt_path}")
            return mapping
        with open(dir_txt_path, "r", encoding="cp1252") as f:
            for line in f:
                line = line.strip()
                if not line.lower().endswith(".png"):
                    continue
                tokens = line.split()
                file_name = tokens[-1]
                if file_name.lower() == "dir.txt":
                    continue
                base_name = os.path.splitext(file_name)[0]
                parts = base_name.split("_")
                filtered_parts = []
                for part in parts:
                    if part.lower() == "portrait" or part.isdigit():
                        continue
                    filtered_parts.append(part)
                if filtered_parts:
                    candidate = " ".join(filtered_parts)
                    normalized_candidate = self.normalize_name(candidate)
                    mapping[normalized_candidate] = file_name
        return mapping

    def associate_npc_portraits(self):
        portrait_mapping = self.build_portrait_mapping()
        if not portrait_mapping:
            print("No portrait mapping was built.")
            return
        db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT Name, Portrait FROM npcs")
        npc_rows = cursor.fetchall()
        modified = False
        for npc in npc_rows:
            npc_name = npc["Name"].strip()
            normalized_npc = self.normalize_name(npc_name)
            if normalized_npc in portrait_mapping:
                portrait_file = portrait_mapping[normalized_npc]
                if not npc["Portrait"] or npc["Portrait"].strip() == "":
                    campaign_dir = ConfigHelper.get_campaign_dir()
                    new_portrait_path = os.path.join(campaign_dir, "assets", "portraits", portrait_file)
                    cursor.execute("UPDATE npcs SET Portrait = ? WHERE Name = ?", (new_portrait_path, npc_name))
                    print(f"Associated portrait '{portrait_file}' with NPC '{npc_name}'")
                    modified = True
        if modified:
            conn.commit()
            print("NPC database updated with associated portraits.")
        else:
            print("No NPC records were updated. Either all have portraits or no matches were found.")
        conn.close()

    def _on_ctrl_f(self, event=None):
        """Global Ctrl+F binding: only opens search when GM screen is active."""
        if self.current_gm_view:
            self.current_gm_view.open_global_search()
         # otherwise ignore silently


    def open_audio_bar(self):
        try:
            window = self.audio_bar_window
        except AttributeError:
            window = None
        try:
            if window is None or not window.winfo_exists():
                self.audio_bar_window = AudioBarWindow(self, controller=self.audio_controller)
                self.audio_bar_window.bind("<Destroy>", self._on_audio_bar_destroyed)
                window = self.audio_bar_window
            window.show()
        except Exception as exc:
            messagebox.showerror("Error", f"Failed to open Audio Controls Bar:\n{exc}")

    def _on_audio_bar_destroyed(self, event=None):
        if event is None or event.widget is self.audio_bar_window:
            self.audio_bar_window = None
            try:
                if self.dice_bar_window and self.dice_bar_window.winfo_exists():
                    self.dice_bar_window._apply_geometry()
            except Exception:
                pass

    def open_dice_bar(self):
        try:
            window = self.dice_bar_window
        except AttributeError:
            window = None
        try:
            if window is None or not window.winfo_exists():
                self.dice_bar_window = DiceBarWindow(self)
                self.dice_bar_window.bind("<Destroy>", self._on_dice_bar_destroyed)
                window = self.dice_bar_window
            window.show()
        except Exception as exc:
            messagebox.showerror("Error", f"Failed to open Dice Bar:\n{exc}")

    def _on_dice_bar_destroyed(self, event=None):
        if event is None:
            self.dice_bar_window = None
            return
        if event.widget is self.dice_bar_window:
            self.dice_bar_window = None

    def _on_system_changed(self, _config) -> None:
        """Refresh open dice windows when the campaign system changes."""

        try:
            dice_markup.invalidate_action_pattern_cache()
        except Exception as exc:
            log_warning(
                f"Failed to invalidate analyzer cache: {exc}",
                func_name="MainWindow._on_system_changed",
            )

        try:
            self.update_sidebar_metadata()
        except Exception as exc:
            log_warning(
                f"Failed to update sidebar metadata: {exc}",
                func_name="MainWindow._on_system_changed",
            )

        try:
            window = getattr(self, "dice_bar_window", None)
            if window is not None and window.winfo_exists():
                window.refresh_system_settings()
        except Exception as exc:
            log_warning(
                f"Failed to refresh dice bar after system change: {exc}",
                func_name="MainWindow._on_system_changed",
            )

        try:
            roller = getattr(self, "dice_roller_window", None)
            if roller is not None and roller.winfo_exists():
                roller.refresh_system_settings()
        except Exception as exc:
            log_warning(
                f"Failed to refresh dice roller after system change: {exc}",
                func_name="MainWindow._on_system_changed",
            )

    def _reload_active_campaign_system(self) -> None:
        """Force the system configuration to match the active campaign DB."""

        try:
            config = system_config.refresh_current_system()
        except Exception as exc:
            log_warning(
                f"Failed to reload campaign system after database switch: {exc}",
                func_name="MainWindow._reload_active_campaign_system",
            )
            return

        if config is None:
            log_warning(
                "No campaign system is configured for the selected database.",
                func_name="MainWindow._reload_active_campaign_system",
            )
        # Apply campaign-specific theme when DB changes
        try:
            from modules.helpers import theme_manager
            theme_key = theme_manager.get_theme()
            theme_manager.apply_theme(theme_key)
            self._on_theme_changed(theme_key)
        except Exception:
            pass

    def destroy(self) -> None:
        listener = getattr(self, "_system_listener_unsub", None)
        if callable(listener):
            try:
                listener()
            except Exception as exc:
                log_warning(
                    f"Failed to unregister system listener: {exc}",
                    func_name="MainWindow.destroy",
                )
            finally:
                self._system_listener_unsub = None
        super().destroy()

    def open_sound_manager(self):
        try:
            window = self.sound_manager_window
        except AttributeError:
            window = None
        try:
            if window is None or not window.winfo_exists():
                self.sound_manager_window = SoundManagerWindow(self, controller=self.audio_controller)
                self.sound_manager_window.bind("<Destroy>", self._on_sound_manager_destroyed)
                window = self.sound_manager_window
            window.show()
        except Exception as exc:
            messagebox.showerror("Error", f"Failed to open Sound & Music Manager:\n{exc}")

    def _on_sound_manager_destroyed(self, event=None):
        if event is None:
            self.sound_manager_window = None
            return
        if event.widget is self.sound_manager_window:
            self.sound_manager_window = None

    def open_dice_roller(self):
        try:
            window = self.dice_roller_window
            if window is None or not window.winfo_exists():
                self.dice_roller_window = DiceRollerWindow(self)
                self.dice_roller_window.bind("<Destroy>", self._on_dice_window_destroyed)
            self.dice_roller_window.show()
        except Exception as exc:
            messagebox.showerror("Error", f"Failed to open Dice Roller:\n{exc}")

    def _on_dice_window_destroyed(self, event=None):
        if event is None:
            self.dice_roller_window = None
            return
        if event.widget is self.dice_roller_window:
            self.dice_roller_window = None

    def open_whiteboard(self):
        log_info("Opening Whiteboard", func_name="main_window.MainWindow.open_whiteboard")
        self.clear_current_content()
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        board_frame = ctk.CTkFrame(container)
        board_frame.pack(fill="both", expand=True)

        self.whiteboard_controller = WhiteboardController(board_frame, root_app=self)

        def _on_destroy(_event=None):
            self._teardown_whiteboard_controller()

        container.bind("<Destroy>", _on_destroy)
        self.current_open_view = container
        self.current_open_entity = None

    def map_tool(self, map_name=None):
        log_info(
            f"Opening Map Tool (map={map_name})",
            func_name="main_window.MainWindow.map_tool",
        )
        # If a GM screen is active, open as a tab instead
        try:
            if getattr(self, "current_gm_view", None) is not None and self.current_gm_view.winfo_exists():
                self.current_gm_view.open_map_tool_tab(map_name)
                return
        except Exception:
            pass

        existing = getattr(self, "_map_tool_window", None)
        if existing is not None and existing.winfo_exists():
            existing.lift()
            existing.focus_force()
            existing.attributes("-topmost", True)
            existing.after_idle(lambda: existing.attributes("-topmost", False))
            controller = getattr(self, "map_controller", None)
            if map_name and controller and hasattr(controller, "open_map_by_name"):
                # Fit to the window when opening from external GM windows
                controller.open_map_by_name(map_name)
            return

        maps_wrapper = GenericModelWrapper("maps")

        top = ctk.CTkToplevel(self)
        top.lift()
        top.focus_force()
        top.attributes("-topmost", True)
        top.after_idle(lambda: top.attributes("-topmost", False))
        top.title("Map Tool")
        top.geometry("1920x1080+0+0")

        map_frame = ctk.CTkFrame(top)
        map_frame.pack(fill="both", expand=True)

        self.map_controller = DisplayMapController(
            map_frame,
            maps_wrapper,
            load_template("maps"),
            root_app=self,
        )
        if map_name and hasattr(self.map_controller, "open_map_by_name"):
            # Fit to the window on initial open
            self.map_controller.open_map_by_name(map_name)

        def _on_close():
            try:
                controller = getattr(self, "map_controller", None)
                if controller is not None:
                    controller.close_web_display()
            except Exception:
                log_exception("Error while closing web display")
            finally:
                self._map_tool_window = None
                self.map_controller = None
                top.destroy()

        top.protocol("WM_DELETE_WINDOW", _on_close)
        self._map_tool_window = top

if __name__ == "__main__":
    if "--apply-update" in sys.argv:
        sys.argv.remove("--apply-update")
        from scripts import apply_update as _apply_update

        _apply_update.main()
    else:
        app = MainWindow()
        app.mainloop()
